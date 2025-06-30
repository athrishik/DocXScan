"""
DocXReplace v3.0 - Ultra Modern Document Replacement Tool
Copyright 2025 Hrishik Kunduru. All rights reserved.

A professional tool for automated document replacement with legal token migration support.
Features include regex patterns, safe backup system, batch processing, and comprehensive logging.
"""

import subprocess
import tkinter as tk
from tkinter import filedialog, messagebox, ttk, Menu
import tempfile
import os
import re
import zipfile
import pandas as pd
from datetime import datetime
from docx import Document
import json
import configparser
import sys
import ctypes
import threading
import shutil
from pathlib import Path
import traceback
import platform
import logging
import time
from typing import Dict, List, Tuple, Optional, Any

# DPI awareness and console hiding for Windows
if platform.system() == "Windows":
    try:
        ctypes.windll.shcore.SetProcessDpiAwareness(2)
        ctypes.windll.user32.ShowWindow(ctypes.windll.kernel32.GetConsoleWindow(), 0)
    except:
        pass

class ModernColors:
    """Premium color palette matching Document Tools Suite exactly"""
    BG_PRIMARY = "#0A0E1A"      # Dark background
    BG_SECONDARY = "#0F1419"     # Darker variation
    
    GLASS_BG = "#1A1D29"        # Glass effect backgrounds
    GLASS_BORDER = "#2A2D3A"    # Glass borders
    GLASS_HOVER = "#1F2235"     # Glass hover state
    
    CARD_BG = "#161925"         # Card backgrounds
    CARD_BORDER = "#252837"     # Card borders
    CARD_HOVER = "#1C1F2E"      # Card hover state
    
    PRIMARY = "#2563EB"         # Blue theme
    PRIMARY_HOVER = "#1D4ED8"   # Darker blue for hover
    SUCCESS = "#10B981"         # Green theme
    SUCCESS_HOVER = "#059669"   # Darker green for hover
    WARNING = "#C56C86"         # Pink tone for warnings
    ERROR = "#FF7582"           # Coral red for errors
    
    TEXT_PRIMARY = "#FFFFFF"    # Pure white text 
    TEXT_SECONDARY = "#E5E7EB"  # Light gray
    TEXT_TERTIARY = "#C5A7A7"   # Warm gray with pink undertone 
    TEXT_MUTED = "#8B7D8B"      # Muted purple-gray 
    
    FOCUS_RING = "#2563EB"      # Blue focus ring
    
    # Gradient colors for decorative elements
    GRADIENT_1 = "#2563EB"      # Blue
    GRADIENT_2 = "#725A7A"      # Purple
    GRADIENT_3 = "#10B981"      # Green
    GRADIENT_4 = "#FF7582"      # Coral

class BlurredBackground(tk.Canvas):
    """Creates a gradient background effect matching the suite"""
    
    def __init__(self, parent, width, height):
        super().__init__(parent, width=width, height=height, highlightthickness=0, bd=0)
        self.configure(bg=ModernColors.BG_PRIMARY)
        self.width = width
        self.height = height
        self.create_gradient_background()
    
    def create_gradient_background(self):
        """Create gradient background with decorative elements"""
        for i in range(self.height):
            ratio = i / self.height
            color_val = int(10 + (25 - 10) * (1 - ratio))
            color = f"#{color_val:02x}{color_val + 2:02x}{color_val + 8:02x}"
            self.create_line(0, i, self.width, i, fill=color, width=1)
        
        # Add decorative circles
        self.create_blur_circle(120, 80, 50)
        self.create_blur_circle(self.width - 100, 150, 40)
        self.create_blur_circle(250, self.height - 120, 60)

    def create_blur_circle(self, x, y, radius):
        """Create decorative circle"""
        colors = [ModernColors.PRIMARY, ModernColors.SUCCESS, ModernColors.WARNING]
        for i, color in enumerate(colors[:2]):
            r = radius - i * 15
            if r > 0:
                self.create_oval(x - r, y - r, x + r, y + r,
                               fill=color, outline="", stipple="gray12")

class ModernCard(tk.Frame):
    """Modern card component"""
    
    def __init__(self, parent, title, width=None, height=None):
        super().__init__(parent)
        self.title = title
        self.font_family = self.get_system_font()
        
        self.setup_card(width, height)
    
    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")
    
    def setup_card(self, width, height):
        """Setup card styling"""
        self.configure(
            bg=ModernColors.CARD_BG,
            relief="flat",
            bd=0,
            highlightthickness=1,
            highlightbackground=ModernColors.CARD_BORDER
        )
        
        if width and height:
            self.configure(width=width, height=height)
            self.pack_propagate(False)
        
        # Add this line at the end:
        self.create_header()
            
    def create_header(self):
        """Create card header with title"""
        if self.title:
            header_frame = tk.Frame(self, bg=ModernColors.CARD_BG)
            header_frame.pack(fill=tk.X, padx=12, pady=(8, 0))
            
            tk.Label(header_frame, text=self.title,
                    font=(self.font_family, 11, "bold"),
                    bg=ModernColors.CARD_BG,
                    fg=ModernColors.TEXT_PRIMARY).pack(anchor="w")
            # Status dot
            tk.Label(header_frame, text="●",
                    font=("Arial", 8),
                    fg=ModernColors.SUCCESS,
                    bg=ModernColors.CARD_BG).pack(side=tk.RIGHT)

    def add_content(self, content_widget):
        """Add content to the card"""
        content_widget.pack(fill=tk.BOTH, expand=True)

class ModernButton(tk.Button):
    """Modern styled button"""
    
    def __init__(self, parent, text, command, style="primary", **kwargs):
        self.font_family = self.get_system_font()
        
        # Style configurations
        styles = {
            "primary": {
                "bg": ModernColors.SUCCESS,  # Changed from PRIMARY to SUCCESS (green)
                "hover_bg": "#059669",  # Darker green hover
                "fg": "white"
            },
            "success": {
                "bg": ModernColors.SUCCESS,
                "hover_bg": "#059669",  # Darker green that matches the suite hover style
                "fg": "white"
            },
            "warning": {
                "bg": ModernColors.WARNING,
                "hover_bg": "#D97706",
                "fg": "white"
            },
            "error": {
                "bg": ModernColors.ERROR,  # Red for stop button
                "hover_bg": "#DC2626",
                "fg": "white"
            },
            "secondary": {
                "bg": ModernColors.GLASS_BG,
                "hover_bg": ModernColors.GLASS_HOVER,
                "fg": ModernColors.TEXT_PRIMARY
            }
        }
        
        config = styles.get(style, styles["primary"])
        
        super().__init__(
            parent,
            text=text,
            command=command,
            font=(self.font_family, 10, "bold"),
            bg=config["bg"],
            fg=config["fg"],
            activebackground=config["bg"],
            activeforeground=config["fg"],
            bd=0,
            relief="flat",
            padx=16,
            pady=8,
            cursor="hand2",
            **kwargs
        )
        
        self.hover_bg = config["hover_bg"]
        self.normal_bg = config["bg"]
        
        self.bind("<Enter>", self.on_enter)
        self.bind("<Leave>", self.on_leave)
    
    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")
    
    def on_enter(self, event):
        self.configure(bg=self.hover_bg)
    
    def on_leave(self, event):
        self.configure(bg=self.normal_bg)

class ModernCombobox(tk.Frame):
    """Modern styled combobox with label - DARK THEME LIKE DOCXSCAN"""
    
    def __init__(self, parent, label, variable=None, values=None, **kwargs):
        super().__init__(parent, bg=ModernColors.BG_PRIMARY)
        self.font_family = self.get_system_font()
        
        # Label
        tk.Label(self, text=label,
                font=(self.font_family, 10, "bold"),
                bg=ModernColors.BG_PRIMARY,
                fg=ModernColors.TEXT_PRIMARY).pack(anchor="w", pady=(0, 3))
        
        # Get the global style and configure it EXACTLY like DocXScan
        style = ttk.Style()
        
        # Configure global options - EXACT COPY from DocXScan
        parent.master.option_add('*TCombobox*Listbox.background', ModernColors.BG_SECONDARY)
        parent.master.option_add('*TCombobox*Listbox.foreground', ModernColors.TEXT_PRIMARY)
        parent.master.option_add('*TCombobox*Listbox.selectBackground', ModernColors.PRIMARY)
        parent.master.option_add('*TCombobox*Listbox.selectForeground', 'white')
        parent.master.option_add('*TCombobox*Listbox.borderWidth', '0')
        parent.master.option_add('*TCombobox*Listbox.relief', 'flat')
        
        # Force override ALL combobox styling - EXACT COPY from DocXScan
        style.configure("Modern.TCombobox",
                       fieldbackground=ModernColors.BG_SECONDARY,
                       background=ModernColors.BG_SECONDARY,
                       foreground=ModernColors.TEXT_PRIMARY,
                       borderwidth=1,
                       relief="flat",
                       selectbackground=ModernColors.PRIMARY,
                       selectforeground="white",
                       arrowcolor=ModernColors.TEXT_PRIMARY,
                       insertcolor=ModernColors.TEXT_PRIMARY,
                       lightcolor=ModernColors.BG_SECONDARY,
                       darkcolor=ModernColors.BG_SECONDARY,
                       bordercolor=ModernColors.CARD_BORDER,
                       focuscolor=ModernColors.PRIMARY)
        
        # Aggressive state mapping to override everything - EXACT COPY from DocXScan
        style.map("Modern.TCombobox",
                 fieldbackground=[("readonly", ModernColors.BG_SECONDARY),
                                ("active", ModernColors.BG_SECONDARY),
                                ("focus", ModernColors.BG_SECONDARY),
                                ("disabled", ModernColors.BG_SECONDARY),
                                ("pressed", ModernColors.BG_SECONDARY),
                                ("!readonly", ModernColors.BG_SECONDARY),
                                ("!focus", ModernColors.BG_SECONDARY),
                                ("!active", ModernColors.BG_SECONDARY)],
                 background=[("readonly", ModernColors.BG_SECONDARY),
                           ("active", ModernColors.BG_SECONDARY),
                           ("focus", ModernColors.BG_SECONDARY),
                           ("disabled", ModernColors.BG_SECONDARY),
                           ("pressed", ModernColors.BG_SECONDARY),
                           ("!readonly", ModernColors.BG_SECONDARY),
                           ("!focus", ModernColors.BG_SECONDARY),
                           ("!active", ModernColors.BG_SECONDARY)],
                 foreground=[("readonly", ModernColors.TEXT_PRIMARY),
                           ("active", ModernColors.TEXT_PRIMARY),
                           ("focus", ModernColors.TEXT_PRIMARY),
                           ("disabled", ModernColors.TEXT_MUTED),
                           ("pressed", ModernColors.TEXT_PRIMARY),
                           ("!readonly", ModernColors.TEXT_PRIMARY)],
                 bordercolor=[("focus", ModernColors.PRIMARY),
                            ("active", ModernColors.CARD_BORDER),
                            ("readonly", ModernColors.CARD_BORDER),
                            ("!focus", ModernColors.CARD_BORDER)],
                 arrowcolor=[("readonly", ModernColors.TEXT_PRIMARY),
                           ("active", ModernColors.TEXT_PRIMARY),
                           ("focus", ModernColors.TEXT_PRIMARY),
                           ("disabled", ModernColors.TEXT_MUTED),
                           ("pressed", ModernColors.TEXT_PRIMARY),
                           ("!readonly", ModernColors.TEXT_PRIMARY)])
        
        self.combobox = ttk.Combobox(self,
                                    textvariable=variable,
                                    values=values or [],
                                    font=(self.font_family, 10),
                                    style="Modern.TCombobox",
                                    state="readonly",
                                    **kwargs)
        self.combobox.pack(fill=tk.X, ipady=3)
        
        # Force update and apply theme immediately
        self.combobox.update()
        self.update_idletasks()
    
    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")

class ModernProgressBar(tk.Frame):
    """Modern progress bar with label"""
    
    def __init__(self, parent):
        super().__init__(parent, bg=ModernColors.BG_PRIMARY)
        self.font_family = self.get_system_font()
        
        # Progress label
        self.progress_label = tk.Label(self, text="Ready to process",
                                      font=(self.font_family, 11, "bold"),
                                      bg=ModernColors.BG_PRIMARY,
                                      fg=ModernColors.TEXT_PRIMARY,
                                      anchor="w")
        self.progress_label.pack(anchor="w", pady=(0, 4))
        
        # Progress bar - thinner
        style = ttk.Style()
        style.configure("Modern.Horizontal.TProgressbar",
                       background=ModernColors.PRIMARY,
                       troughcolor=ModernColors.CARD_BG,
                       borderwidth=0,
                       lightcolor=ModernColors.PRIMARY,
                       darkcolor=ModernColors.PRIMARY)
        
        self.progress_bar = ttk.Progressbar(self,
                                           style="Modern.Horizontal.TProgressbar",
                                           mode="determinate",
                                           maximum=100)
        self.progress_bar.pack(fill=tk.X, ipady=2)  # Reduced from 6 to 2
    
    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")
    
    def set_progress(self, value, text=None):
        self.progress_bar["value"] = value
        if text:
            self.progress_label.configure(text=text)

class ModernTextArea(tk.Frame):
    """Modern text area with scrollbar"""
    
    def __init__(self, parent, height=10):
        super().__init__(parent, bg=ModernColors.BG_PRIMARY)
        self.font_family = self.get_system_font()
        
        # Create text widget with scrollbar
        text_frame = tk.Frame(self, bg=ModernColors.CARD_BG)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        self.text_widget = tk.Text(text_frame,
                                  height=height,
                                  font=(self.font_family, 10),
                                  bg=ModernColors.CARD_BG,
                                  fg=ModernColors.TEXT_PRIMARY,
                                  insertbackground=ModernColors.TEXT_PRIMARY,
                                  relief="flat",
                                  bd=0,
                                  wrap=tk.WORD,
                                  padx=12,
                                  pady=8)
        
        scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=self.text_widget.yview)
        self.text_widget.configure(yscrollcommand=scrollbar.set)
        
        self.text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    
    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")
    
    def insert(self, text):
        self.text_widget.insert(tk.END, text)
        self.text_widget.see(tk.END)
    
    def clear(self):
        self.text_widget.delete(1.0, tk.END)

class Logger:
    """Professional logging system"""

    def __init__(self):
        # Create logs directory
        Path("logs").mkdir(exist_ok=True)

        # Setup logging
        self.logger = logging.getLogger("DocXReplace")
        self.logger.setLevel(logging.INFO)

        # File handler
        file_handler = logging.FileHandler(f'logs/docx_replace_{datetime.now().strftime("%Y%m%d")}.log')
        file_handler.setLevel(logging.INFO)

        # Console handler
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.WARNING)

        # Formatter
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)

        self.logger.addHandler(file_handler)
        self.logger.addHandler(console_handler)

    def info(self, message: str):
        self.logger.info(message)

    def warning(self, message: str):
        self.logger.warning(message)

    def error(self, message: str):
        self.logger.error(message)

    def debug(self, message: str):
        self.logger.debug(message)

class SettingsManager:
    """Manage application settings and configuration"""

    def __init__(self, config_file: str = "docxreplace.ini"):
        self.config_file = config_file
        self.config = configparser.ConfigParser()
        self.default_settings = {
            "theme": "dark",
            "auto_backup": "True",
            "max_recent_files": "10",
            "default_output_folder": str(Path.home() / "Documents" / "DocReplace_Output"),
            "log_level": "INFO",
            "auto_save_config": "True"
        }

    def load_config(self) -> Dict[str, str]:
        """Load configuration from file"""
        if os.path.exists(self.config_file):
            try:
                self.config.read(self.config_file)
                return dict(self.config["LAST"]) if "LAST" in self.config else {}
            except Exception:
                return {}
        return {}

    def save_config(self, settings: Dict[str, str]):
        """Save configuration to file"""
        try:
            if "LAST" not in self.config:
                self.config["LAST"] = {}

            for key, value in settings.items():
                self.config["LAST"][key] = str(value)

            with open(self.config_file, "w") as f:
                self.config.write(f)
        except Exception as e:
            print(f"Error saving config: {e}")

    def get_setting(self, key: str, fallback: str = "") -> str:
        """Get a specific setting"""
        return self.config.get("LAST", key, fallback=fallback)

class DocumentProcessor:
    """Handle document processing operations"""

    def __init__(self, logger: Logger):
        self.logger = logger

    def validate_replacement_map(self, replacement_map: Dict[str, str]) -> List[str]:
        """Validate replacement map for common issues"""
        errors = []

        for old_text, new_text in replacement_map.items():
            if not old_text.strip():
                errors.append("Empty search pattern found")
            if len(old_text) > 1000:
                errors.append(f"Search pattern too long: {old_text[:50]}...")
            if not isinstance(new_text, str):
                errors.append(f"Invalid replacement type for '{old_text}': expected string")

        return errors

    def perform_replacement_in_doc(self, doc: Document, file_path: str, 
                                 replacement_map: Dict[str, str], 
                                 regex_mode: bool = False) -> Tuple[int, List[Dict]]:
        """Perform replacements in a document with enhanced error handling"""
        replacements_made = 0
        replacement_details = []

        try:
            # Process paragraphs
            for para_idx, para in enumerate(doc.paragraphs):
                original_text = para.text
                modified_text = original_text

                for old_text, new_text in replacement_map.items():
                    try:
                        if regex_mode:
                            if "{{match}}" in new_text:
                                def replace_with_match(match):
                                    matched_groups = match.groups()
                                    replacement = new_text
                                    if matched_groups:
                                        for i, group in enumerate(matched_groups):
                                            replacement = replacement.replace("{{match}}", group, 1)
                                    else:
                                        replacement = replacement.replace("{{match}}", match.group(0))
                                    return replacement
                                modified_text = re.sub(old_text, replace_with_match, modified_text)
                            else:
                                modified_text = re.sub(old_text, new_text, modified_text)
                        else:
                            if old_text in modified_text:
                                modified_text = modified_text.replace(old_text, new_text)
                    except re.error as e:
                        self.logger.warning(f"Invalid regex pattern '{old_text}': {e}")
                        continue
                    except Exception as e:
                        self.logger.error(f"Error processing pattern '{old_text}': {e}")
                        continue

                if modified_text != original_text:
                    try:
                        para.clear()
                        para.add_run(modified_text)
                    except Exception:
                        para.text = modified_text

                    replacements_made += 1
                    replacement_details.append({
                        'location': f'paragraph_{para_idx}',
                        'original': original_text[:100] + '...' if len(original_text) > 100 else original_text,
                        'modified': modified_text[:100] + '...' if len(modified_text) > 100 else modified_text
                    })

            # Process tables
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            original_text = para.text
                            modified_text = original_text

                            for old_text, new_text in replacement_map.items():
                                try:
                                    if regex_mode:
                                        if "{{match}}" in new_text:
                                            def replace_with_match(match):
                                                matched_groups = match.groups()
                                                replacement = new_text
                                                if matched_groups:
                                                    for i, group in enumerate(matched_groups):
                                                        replacement = replacement.replace("{{match}}", group, 1)
                                                else:
                                                    replacement = replacement.replace("{{match}}", match.group(0))
                                                return replacement
                                            modified_text = re.sub(old_text, replace_with_match, modified_text)
                                        else:
                                            modified_text = re.sub(old_text, new_text, modified_text)
                                    else:
                                        if old_text in modified_text:
                                            modified_text = modified_text.replace(old_text, new_text)
                                except re.error:
                                    continue
                                except Exception as e:
                                    self.logger.error(f"Error in table processing: {e}")
                                    continue

                            if modified_text != original_text:
                                try:
                                    para.clear()
                                    para.add_run(modified_text)
                                except Exception:
                                    para.text = modified_text

                                replacements_made += 1
                                replacement_details.append({
                                    'location': f'table_{table_idx}_row_{row_idx}_cell_{cell_idx}_para_{para_idx}',
                                    'original': original_text[:50] + '...' if len(original_text) > 50 else original_text,
                                    'modified': modified_text[:50] + '...' if len(modified_text) > 50 else modified_text
                                })

        except Exception as e:
            self.logger.error(f"Critical error processing document {file_path}: {e}")
            raise

        return replacements_made, replacement_details

class ToolTip:
    """Add tooltips to widgets"""

    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tooltip_window = None
        self.widget.bind("<Enter>", self.on_enter)
        self.widget.bind("<Leave>", self.on_leave)

    def on_enter(self, event):
        try:
            x, y, _, _ = self.widget.bbox("insert") if hasattr(self.widget, 'bbox') else (0, 0, 0, 0)
        except Exception:
            x, y = 0, 0
        x += self.widget.winfo_rootx() + 25
        y += self.widget.winfo_rooty() + 25

        self.tooltip_window = tk.Toplevel(self.widget)
        self.tooltip_window.wm_overrideredirect(True)
        self.tooltip_window.wm_geometry(f"+{x}+{y}")

        label = tk.Label(self.tooltip_window, text=self.text, 
                        background=ModernColors.CARD_BG, foreground=ModernColors.TEXT_PRIMARY,
                        relief="solid", borderwidth=1, font=("Arial", 9))
        label.pack()

    def on_leave(self, event):
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None

class DocXReplaceApp:
    """Main application class with modern UI"""

    def __init__(self):
        # Initialize components
        self.logger = Logger()
        self.settings = SettingsManager()
        self.document_processor = DocumentProcessor(self.logger)

        # State variables
        self.replacement_map = {}
        self.loaded_files = []
        self.backup_history = []
        self.temp_directories = []
        self.files_source_type = "folder"
        self.stop_processing = False
        self.recent_files = []
        self.recent_replacements = []

        # Font
        self.font_family = self.get_system_font()

        # Setup GUI
        self.setup_main_window()
        self.setup_variables()
        self.setup_ui()
        self.setup_menu()
        self.setup_shortcuts()

        # Load configuration
        self.load_configuration()

        self.logger.info("DocXReplace v3.0 started successfully")

    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")

    def setup_main_window(self):
        """Initialize the main window - EXACT SIZING MATCH to DocXScan"""
        self.root = tk.Tk()
        self.root.title("DocXReplace v3.0 - Professional Document Replacement Tool")
        self.root.geometry("1300x800")
        self.root.configure(bg=ModernColors.BG_PRIMARY)
        self.root.resizable(True, True)
        self.root.minsize(1100, 700)
        # Configure global ttk theme for dark mode 
        self.configure_dark_theme()
        

    def configure_dark_theme(self):
        """Configure dark theme for all ttk widgets - EXACT COPY from DocXScan"""
        style = ttk.Style()
        
        # Force a dark theme
        try:
            style.theme_use('clam')  # Use clam as base for better customization
        except:
            pass
        
        # Configure global options for better dark theme support - EXACT COPY from DocXScan
        self.root.option_add('*TCombobox*Listbox.background', ModernColors.BG_SECONDARY)
        self.root.option_add('*TCombobox*Listbox.foreground', ModernColors.TEXT_PRIMARY)
        self.root.option_add('*TCombobox*Listbox.selectBackground', ModernColors.PRIMARY)
        self.root.option_add('*TCombobox*Listbox.selectForeground', 'white')
        self.root.option_add('*TCombobox*Listbox.borderWidth', '0')
        self.root.option_add('*TCombobox*Listbox.relief', 'flat')
        
        # Configure ALL TCombobox styles aggressively - EXACT COPY from DocXScan
        style.configure("TCombobox",
                       fieldbackground=ModernColors.BG_SECONDARY,
                       background=ModernColors.BG_SECONDARY,
                       foreground=ModernColors.TEXT_PRIMARY,
                       borderwidth=1,
                       relief="flat",
                       selectbackground=ModernColors.PRIMARY,
                       selectforeground="white",
                       arrowcolor=ModernColors.TEXT_PRIMARY,
                       insertcolor=ModernColors.TEXT_PRIMARY,
                       lightcolor=ModernColors.BG_SECONDARY,
                       darkcolor=ModernColors.BG_SECONDARY,
                       bordercolor=ModernColors.CARD_BORDER,
                       focuscolor=ModernColors.PRIMARY)
        
        # Map ALL possible states for TCombobox - EXACT COPY from DocXScan
        style.map("TCombobox",
                 fieldbackground=[("readonly", ModernColors.BG_SECONDARY),
                                ("active", ModernColors.BG_SECONDARY),
                                ("focus", ModernColors.BG_SECONDARY),
                                ("disabled", ModernColors.BG_SECONDARY),
                                ("pressed", ModernColors.BG_SECONDARY),
                                ("!readonly", ModernColors.BG_SECONDARY)],
                 background=[("readonly", ModernColors.BG_SECONDARY),
                           ("active", ModernColors.BG_SECONDARY),
                           ("focus", ModernColors.BG_SECONDARY),
                           ("disabled", ModernColors.BG_SECONDARY),
                           ("pressed", ModernColors.BG_SECONDARY),
                           ("!readonly", ModernColors.BG_SECONDARY)],
                 foreground=[("readonly", ModernColors.TEXT_PRIMARY),
                           ("active", ModernColors.TEXT_PRIMARY),
                           ("focus", ModernColors.TEXT_PRIMARY),
                           ("disabled", ModernColors.TEXT_MUTED),
                           ("pressed", ModernColors.TEXT_PRIMARY),
                           ("!readonly", ModernColors.TEXT_PRIMARY)],
                 bordercolor=[("focus", ModernColors.PRIMARY),
                            ("active", ModernColors.PRIMARY),
                            ("readonly", ModernColors.CARD_BORDER),
                            ("disabled", ModernColors.CARD_BORDER),
                            ("pressed", ModernColors.PRIMARY),
                            ("!readonly", ModernColors.CARD_BORDER)],
                 arrowcolor=[("readonly", ModernColors.TEXT_PRIMARY),
                           ("active", ModernColors.TEXT_PRIMARY),
                           ("focus", ModernColors.TEXT_PRIMARY),
                           ("disabled", ModernColors.TEXT_MUTED),
                           ("pressed", ModernColors.TEXT_PRIMARY),
                           ("!readonly", ModernColors.TEXT_PRIMARY)])
        
        # Configure scrollbar for dark theme - EXACT COPY from DocXScan
        style.configure("Vertical.TScrollbar",
                       background=ModernColors.CARD_BG,
                       troughcolor=ModernColors.BG_SECONDARY,
                       bordercolor=ModernColors.CARD_BORDER,
                       arrowcolor=ModernColors.TEXT_PRIMARY,
                       darkcolor=ModernColors.CARD_BG,
                       lightcolor=ModernColors.CARD_BG)
        
        # Map scrollbar states
        style.map("Vertical.TScrollbar",
                 background=[("active", ModernColors.CARD_HOVER),
                           ("pressed", ModernColors.PRIMARY)])
        
        # Configure progress bar theme with your colors
        style.configure("Modern.Horizontal.TProgressbar",
                       background=ModernColors.SUCCESS,  # Use your green
                       troughcolor=ModernColors.CARD_BG,
                       borderwidth=0,
                       lightcolor=ModernColors.SUCCESS,
                       darkcolor=ModernColors.SUCCESS)
        
        # Also configure any other potential ttk widgets
        style.configure("TEntry",
                       fieldbackground=ModernColors.BG_SECONDARY,
                       background=ModernColors.BG_SECONDARY,
                       foreground=ModernColors.TEXT_PRIMARY,
                       bordercolor=ModernColors.CARD_BORDER,
                       insertcolor=ModernColors.TEXT_PRIMARY)
        
        style.map("TEntry",
                 fieldbackground=[("focus", ModernColors.BG_SECONDARY),
                                ("active", ModernColors.BG_SECONDARY)],
                 bordercolor=[("focus", ModernColors.PRIMARY),
                            ("active", ModernColors.PRIMARY)])
    

    def setup_variables(self):
        """Initialize tkinter variables"""
        self.selected_files_folder = tk.StringVar()
        self.replacement_file_path = tk.StringVar()
        self.backup_folder = tk.StringVar()
        self.replacement_mode = tk.StringVar(value="Create Modified Copies (originals untouched)")
        self.loaded_replacement_file = tk.StringVar(value="No replacement file loaded.")
        self.loaded_files_status = tk.StringVar(value="No files loaded.")
        self.excel_file_path = tk.StringVar()
        self.regex_mode = tk.BooleanVar(value=False)
        self.auto_backup = tk.BooleanVar(value=True)
        self.status_text = tk.StringVar(value="Ready")
        self.progress_text = tk.StringVar(value="")
        self.eta_text = tk.StringVar(value="")
        
    def setup_menu(self):
        """Setup application menu"""
        menubar = Menu(self.root)

        # File menu
        file_menu = Menu(menubar, tearoff=0)
        file_menu.add_command(label="Load Files...", command=self.browse_files_folder, accelerator="Ctrl+O")
        file_menu.add_command(label="Load Replacements...", command=self.load_replacement_file, accelerator="Ctrl+R")
        file_menu.add_separator()

        # Recent files submenu
        self.recent_menu = Menu(file_menu, tearoff=0)
        file_menu.add_cascade(label="Recent Files", menu=self.recent_menu)

        file_menu.add_separator()
        file_menu.add_command(label="Save Configuration", command=self.save_batch_config)
        file_menu.add_command(label="Load Configuration", command=self.load_batch_config)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.on_closing)

        # Tools menu
        tools_menu = Menu(menubar, tearoff=0)
        tools_menu.add_command(label="Create Template", command=self.create_replacement_template)
        tools_menu.add_command(label="Validate Patterns", command=self.validate_patterns)
        tools_menu.add_command(label="Clean Temp Files", command=self.cleanup_temp_files)
        tools_menu.add_separator()
        tools_menu.add_command(label="Settings", command=self.show_settings)

        # Help menu
        help_menu = Menu(menubar, tearoff=0)
        help_menu.add_command(label="Token Help", command=self.show_token_help)
        help_menu.add_command(label="Regex Help", command=self.show_regex_help)
        help_menu.add_separator()
        help_menu.add_command(label="View Logs", command=self.view_logs)
        help_menu.add_command(label="About", command=self.show_about)

        menubar.add_cascade(label="File", menu=file_menu)
        menubar.add_cascade(label="Tools", menu=tools_menu)
        menubar.add_cascade(label="Help", menu=help_menu)

        self.root.config(menu=menubar)

    def setup_shortcuts(self):
        """Setup keyboard shortcuts"""
        self.root.bind('<Control-o>', lambda e: self.browse_files_folder())
        self.root.bind('<Control-r>', lambda e: self.load_replacement_file())
        self.root.bind('<F5>', lambda e: self.run_replacement_threaded())
        self.root.bind('<Escape>', lambda e: self.stop_replacement())
        self.root.bind('<Control-s>', lambda e: self.save_batch_config())
        self.root.bind('<Control-l>', lambda e: self.load_batch_config())


    def setup_ui(self):
        """Setup the modern user interface"""
        # Background
        bg_canvas = BlurredBackground(self.root, 1300, 800)
        bg_canvas.pack(fill=tk.BOTH, expand=True)
        
        # Main overlay
        main_overlay = tk.Frame(bg_canvas, bg=ModernColors.BG_PRIMARY)
        main_overlay.place(relx=0, rely=0, relwidth=1, relheight=1)
        
        self.create_header(main_overlay)
        self.create_main_content(main_overlay)
        self.create_footer(main_overlay)

    def create_header(self, parent):
        header_frame = tk.Frame(parent, bg=ModernColors.BG_PRIMARY)  # Remove height=70
        header_frame.pack(fill=tk.X, padx=30, pady=(20, 10))
        # Remove header_frame.pack_propagate(False)
            
        # Title
        tk.Label(header_frame,text="DocXReplace v3.0",font=(self.font_family, 24, "bold"),
                 bg=ModernColors.BG_PRIMARY,fg=ModernColors.TEXT_PRIMARY).pack(anchor="w")
        
        tk.Label(header_frame, text="Professional document replacement with legal token migration support",
                font=(self.font_family, 11,"normal"),
                bg=ModernColors.BG_PRIMARY,fg=ModernColors.TEXT_SECONDARY).pack(anchor="w", pady=(2, 0))
        
    def create_main_content(self, parent):
        """Create main content area"""
        content_frame = tk.Frame(parent, bg=ModernColors.BG_PRIMARY)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=(0, 6))  # Reduced padding
        
        # Left column - Configuration (45% width, reduced from 50%)
        left_frame = tk.Frame(content_frame, bg=ModernColors.BG_PRIMARY)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, padx=(0, 10))  # Reduced from 12 to 10
        left_frame.configure(width=int(1300 * 0.45))  # Set explicit width ratio
        
        # Right column - Console (55% width, increased from 50%)
        right_frame = tk.Frame(content_frame, bg=ModernColors.BG_PRIMARY)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        self.create_config_section(left_frame)
        self.create_console_section(right_frame)

    def create_config_section(self, parent):
        """Create configuration section"""
        # File Selection Card - make smaller
        file_card = ModernCard(parent, "📁  File Selection")
        file_card.pack(fill=tk.X, pady=(0, 5))  
        
        file_content = tk.Frame(file_card, bg=ModernColors.CARD_BG)
        file_content.pack(fill=tk.BOTH, expand=True, padx=12, pady=(0, 5))

        # File source buttons
        file_buttons_frame = tk.Frame(file_content, bg=ModernColors.CARD_BG)
        file_buttons_frame.pack(fill=tk.X, pady=(0, 4))  # Reduced from 6 to 4
        
        # Primary button - DocXScan output
        excel_btn = ModernButton(file_buttons_frame, "📊 Load DocXScan Output", self.browse_excel_file)
        excel_btn.pack(side=tk.LEFT, padx=(0, 4))  # Keep at 4
        ToolTip(excel_btn, "Load files from DocXScan Excel report or ZIP archive")
        
        # Secondary buttons
        browse_btn = ModernButton(file_buttons_frame, "📁 Browse Folder", self.browse_files_folder, style="secondary")
        browse_btn.pack(side=tk.LEFT, padx=(0, 4))  # Keep at 4
        ToolTip(browse_btn, "Select a folder containing .docx files")
        
        clean_btn = ModernButton(file_buttons_frame, "🧹 Clean Temp", self.cleanup_temp_files, style="secondary")
        clean_btn.pack(side=tk.LEFT)
        ToolTip(clean_btn, "Clean up temporary extracted files")
        
        # File status
        self.files_status_label = tk.Label(file_content, textvariable=self.loaded_files_status,
                                          font=(self.font_family, 10),
                                          bg=ModernColors.CARD_BG,
                                          fg=ModernColors.TEXT_TERTIARY,
                                          anchor="w")
        self.files_status_label.pack(fill=tk.X, pady=(1, 0))  # Reduced from 2 to 1
        
        # Replacement Configuration Card - make smaller
        repl_card = ModernCard(parent, "🔄  Replacement Configuration")
        repl_card.pack(fill=tk.X, pady=(0, 5))  # Reduced from 6 to 5
        
        repl_content = tk.Frame(repl_card, bg=ModernColors.CARD_BG)
        repl_content.pack(fill=tk.BOTH, expand=True, padx=12, pady=(0, 5))  # Reduced from 6 to 5

        # Replacement file buttons
        repl_buttons_frame = tk.Frame(repl_content, bg=ModernColors.CARD_BG)
        repl_buttons_frame.pack(fill=tk.X, pady=(0, 4))  # Reduced from 6 to 4
        
        load_repl_btn = ModernButton(repl_buttons_frame, "📋 Load Replacements", self.load_replacement_file)
        load_repl_btn.pack(side=tk.LEFT, padx=(0, 4))  # Keep at 4
        ToolTip(load_repl_btn, "Load replacement patterns from JSON file")
        
        template_btn = ModernButton(repl_buttons_frame, "📄 Create Template", self.create_replacement_template, style="secondary")
        template_btn.pack(side=tk.LEFT)
        ToolTip(template_btn, "Create a template replacement file")
        
        # Replacement status
        self.repl_status_label = tk.Label(repl_content, textvariable=self.loaded_replacement_file,
                                         font=(self.font_family, 10),
                                         bg=ModernColors.CARD_BG,
                                         fg=ModernColors.TEXT_TERTIARY,
                                         anchor="w")
        self.repl_status_label.pack(fill=tk.X, pady=(1, 4))  # Reduced top from 2 to 1, bottom from 6 to 4

        # Regex mode and help
        regex_frame = tk.Frame(repl_content, bg=ModernColors.CARD_BG)
        regex_frame.pack(fill=tk.X, pady=(0, 4))  # Reduced from 6 to 4
        
        # Regex checkbox
        self.regex_check = tk.Checkbutton(regex_frame, text="Enable Regex Mode", 
                                         variable=self.regex_mode,
                                         command=self.on_regex_mode_changed,
                                         font=(self.font_family, 10, "bold"),
                                         bg=ModernColors.CARD_BG,
                                         fg=ModernColors.TEXT_PRIMARY,
                                         selectcolor=ModernColors.BG_SECONDARY,
                                         activebackground=ModernColors.CARD_BG,
                                         activeforeground=ModernColors.TEXT_PRIMARY)
        self.regex_check.pack(side=tk.LEFT)
        ToolTip(self.regex_check, "Enable regular expression pattern matching")
        
        # Help buttons
        token_help_btn = ModernButton(regex_frame, "📖 Token Help", self.show_token_help, style="secondary")
        token_help_btn.configure(padx=8, pady=4, font=(self.font_family, 9, "bold"))  # Keep compact
        token_help_btn.pack(side=tk.RIGHT, padx=(3, 0))  # Keep at 3
        
        regex_help_btn = ModernButton(regex_frame, "🔧 Regex Help", self.show_regex_help, style="secondary")
        regex_help_btn.configure(padx=8, pady=4, font=(self.font_family, 9, "bold"))  # Keep compact
        regex_help_btn.pack(side=tk.RIGHT, padx=(3, 0))  # Keep at 3
        
        # Processing Options Card - EXPANDED VERSION
        proc_card = ModernCard(parent, "⚙️  Processing Options")
        proc_card.pack(fill=tk.X, pady=(0, 5))
        
        proc_content = tk.Frame(proc_card, bg=ModernColors.CARD_BG)
        proc_content.pack(fill=tk.X, padx=12, pady=(3, 8))  # Increased bottom padding from 3 to 8
        
        mode_combo = ModernCombobox(proc_content, "Mode:", self.replacement_mode,
                                   values=["Dry Run (preview only)",
                                          "Create Modified Copies (originals untouched)",
                                          "In-place Replace (modify originals)"])
        mode_combo.pack(fill=tk.X, pady=(0, 6))  # Increased from 2 to 6
        
        # Output folder section with more spacing
        output_frame = tk.Frame(proc_content, bg=ModernColors.CARD_BG)
        output_frame.pack(fill=tk.X, pady=(0, 4))
        
        output_btn = ModernButton(output_frame, "📂 Browse Output Folder", self.browse_backup_folder, style="secondary")
        output_btn.configure(pady=6, font=(self.font_family, 9, "bold"))  # Increased pady from 4 to 6
        output_btn.pack(anchor="w")
        ToolTip(output_btn, "Select folder for output files")
        
        # Controls Card - COMPACT VERSION
        controls_card = ModernCard(parent, "🚀  Process Controls")
        controls_card.pack(fill=tk.X)
        
        controls_content = tk.Frame(controls_card, bg=ModernColors.CARD_BG)
        controls_content.pack(fill=tk.X, padx=12, pady=(0, 8))
        
        # Progress bar
        self.progress_bar = ModernProgressBar(controls_content)
        self.progress_bar.pack(fill=tk.X, pady=(0, 6))
        
        # Control buttons frame
        button_frame = tk.Frame(controls_content, bg=ModernColors.CARD_BG)
        button_frame.pack(fill=tk.X)
        
        # All buttons in one row first
        buttons_row = tk.Frame(button_frame, bg=ModernColors.CARD_BG)
        buttons_row.pack(fill=tk.X)
        
        self.start_button = ModernButton(buttons_row, "🚀 Start Replacement", self.run_replacement_threaded)
        self.start_button.configure(width=20)
        self.start_button.pack(side=tk.LEFT)
        
        self.stop_button = ModernButton(buttons_row, "🛑 Stop", self.stop_replacement, style="error") 
        self.stop_button.configure(state="disabled")
        self.stop_button.pack(side=tk.LEFT, padx=(4, 6))
        
        self.undo_button = ModernButton(buttons_row, "↩️ Undo Last", self.undo_last_operation, style="secondary")
        self.undo_button.configure(state="disabled")
        self.undo_button.pack(side=tk.LEFT, padx=(0, 6))
        
        history_btn = ModernButton(buttons_row, "📋 History", self.show_backup_history, style="secondary")
        history_btn.pack(side=tk.LEFT)
        
        # ETA display in separate row below, aligned under start button
        eta_row = tk.Frame(button_frame, bg=ModernColors.CARD_BG)
        eta_row.pack(fill=tk.X, pady=(4, 0))
        
        self.eta_label = tk.Label(eta_row, 
                                 textvariable=self.eta_text,
                                 font=(self.font_family, 9, "normal"),
                                 bg=ModernColors.CARD_BG,
                                 fg=ModernColors.TEXT_TERTIARY,
                                 anchor="w")
        self.eta_label.pack(anchor="w")  # Left-aligned under start button

    def create_console_section(self, parent):
        console_card = ModernCard(parent, "📋  Console Output")
        console_card.pack(fill=tk.BOTH, expand=True)
        # Add clear button to the existing header
        # Find the header frame (first child should be the header)
        header_frame = None
        for child in console_card.winfo_children():
            if isinstance(child, tk.Frame):
                header_frame = child
                break
        
        if header_frame:
            # Add clear button to the right side of the header, before the status dot
            clear_btn = ModernButton(header_frame, "🧹 Clear", self.clear_console, style="secondary")
            clear_btn.configure(padx=8, pady=4, font=(self.font_family, 9, "bold"))
            clear_btn.pack(side=tk.RIGHT, padx=(0, 8))  # Add some padding before the status dot
        console_content = tk.Frame(console_card, bg=ModernColors.CARD_BG)
        console_content.pack(fill=tk.BOTH, expand=True, padx=16, pady=(0, 12))
        
        # Console text area (no separate header needed now)
        self.console = ModernTextArea(console_content, height=16) 
        self.console.pack(fill=tk.BOTH, expand=True)


    
    def create_footer(self, parent):
        """Create footer"""
        footer_frame = tk.Frame(parent, bg=ModernColors.BG_PRIMARY)
        footer_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=30, pady=(0, 10))
        
        tk.Label(footer_frame,
                text="© 2025 Hrishik Kunduru • DocXReplace v3.0 Professional • All Rights Reserved",
                font=(self.font_family, 8, "normal"),
                bg=ModernColors.BG_PRIMARY,
                fg=ModernColors.TEXT_MUTED).pack()

    def center_window(self):
        """Center window on screen"""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')

    def log(self, msg: str):
        """Enhanced logging with timestamp"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        formatted_msg = f"[{timestamp}] {msg}"
        self.console.insert(formatted_msg + "\n")

        # Also log to file
        if "❌" in msg or "Error" in msg:
            self.logger.error(msg)
        elif "⚠️" in msg or "Warning" in msg:
            self.logger.warning(msg)
        else:
            self.logger.info(msg)

    def clear_console(self):
        """Clear console with confirmation for important content"""
        if len(self.console.text_widget.get(1.0, tk.END).strip()) > 100:
            if not messagebox.askyesno("Clear Console", "Clear console output?\n\nThis action cannot be undone."):
                return

        self.console.clear()
        self.log("Console cleared")

    def update_status(self, message: str, is_error: bool = False):
        """Update status message"""
        self.status_text.set(f"{'ERROR: ' if is_error else ''}{message}")
        if is_error:
            self.logger.error(message)
        else:
            self.logger.info(message)

    def update_progress(self, current: int, total: int, eta_seconds: float = 0, current_file: str = ""):
        """Update progress information with proper ETA calculation"""
        percentage = (current / total) * 100 if total > 0 else 0
        
        # Format ETA for standalone display
        if eta_seconds > 60:
            minutes = int(eta_seconds // 60)
            seconds = int(eta_seconds % 60)
            eta_display = f"⏱️ ETA: {minutes}m {seconds}s"
        elif eta_seconds > 0:
            eta_display = f"⏱️ ETA: {int(eta_seconds)}s"
        else:
            eta_display = ""
        
        # Update standalone ETA display
        self.eta_text.set(eta_display)
        
        # Format progress text (without ETA since it's shown separately now)
        if current_file:
            filename = os.path.basename(current_file)
            if len(filename) > 30:  # Increased since ETA is separate
                filename = filename[:27] + "..."
            progress_text = f"{current}/{total} - {filename}"
        else:
            progress_text = f"{current}/{total} files"
        
        # Update progress bar
        self.progress_bar.set_progress(percentage, progress_text)

    def stop_replacement(self):
        """Stop the current replacement process"""
        self.stop_processing = True
        self.eta_text.set("")  # Clear ETA display
        self.log("🛑 Stop requested - finishing current file and stopping...")
        self.update_status("Stopping process...")

    def browse_files_folder(self):
        """Browse for files folder with enhanced feedback"""
        initial_dir = self.get_smart_default_folder()
        path = filedialog.askdirectory(title="Select folder containing files to process", 
                                      initialdir=initial_dir)
        if path:
            self.selected_files_folder.set(path)
            self.files_source_type = "folder"
            self.load_files_from_folder(path)
            self.save_configuration()
            self.add_to_recent_files(path)

    def get_smart_default_folder(self):
        """Get smart default folder based on recent usage"""
        recent_folders = []
        for file_path in self.recent_files:
            if os.path.isdir(file_path):
                recent_folders.append(file_path)
            else:
                parent = os.path.dirname(file_path)
                if os.path.exists(parent):
                    recent_folders.append(parent)
        
        if recent_folders:
            return recent_folders[0]
        
        default_folder = self.settings.get_setting("default_folder", "")
        if default_folder and os.path.exists(default_folder):
            return default_folder
        
        return str(Path.home() / "Documents")

    def browse_backup_folder(self):
        """Browse for backup folder"""
        initial_dir = self.backup_folder.get() or self.settings.get_setting("default_output_folder")
        path = filedialog.askdirectory(title="Select output destination folder", initialdir=initial_dir)
        if path:
            self.backup_folder.set(path)
            self.save_configuration()
            self.update_status(f"Output folder set: {os.path.basename(path)}")

    def browse_excel_file(self):
        """Browse for Excel or ZIP file"""
        path = filedialog.askopenfilename(
            title="Select Excel file from DocXScan or ZIP file",
            filetypes=[("ZIP Files", "*.zip"), ("Excel Files", "*.xlsx"), ("All Files", "*.*")]
        )
        if path:
            if path.lower().endswith('.zip'):
                self.excel_file_path.set(path)
                self.files_source_type = "zip"
                self.load_files_from_zip(path)
            else:
                self.excel_file_path.set(path)
                self.files_source_type = "excel"
                self.load_files_from_excel(path)
            self.add_to_recent_files(path)

    def add_to_recent_files(self, file_path: str):
        """Add file to recent files list"""
        if file_path in self.recent_files:
            self.recent_files.remove(file_path)
        self.recent_files.insert(0, file_path)
        self.recent_files = self.recent_files[:10]
        self.update_recent_menu()

    def update_recent_menu(self):
        """Update recent files menu"""
        self.recent_menu.delete(0, tk.END)

        for i, file_path in enumerate(self.recent_files):
            if os.path.exists(file_path):
                display_name = f"{i+1}. {os.path.basename(file_path)}"
                self.recent_menu.add_command(
                    label=display_name,
                    command=lambda f=file_path: self.load_recent_file(f)
                )

        if not self.recent_files:
            self.recent_menu.add_command(label="No recent files", state="disabled")

    def load_recent_file(self, file_path: str):
        """Load a recent file"""
        if not os.path.exists(file_path):
            messagebox.showerror("File Not Found", f"File not found:\n{file_path}")
            self.recent_files.remove(file_path)
            self.update_recent_menu()
            return

        if file_path.lower().endswith('.zip'):
            self.excel_file_path.set(file_path)
            self.files_source_type = "zip"
            self.load_files_from_zip(file_path)
        elif file_path.lower().endswith('.xlsx'):
            self.excel_file_path.set(file_path)
            self.files_source_type = "excel"
            self.load_files_from_excel(file_path)
        else:
            self.selected_files_folder.set(file_path)
            self.files_source_type = "folder"
            self.load_files_from_folder(file_path)

    def load_files_from_folder(self, folder_path: str):
        """Load files from folder with file counting and filtering"""
        self.loaded_files = []
        if not os.path.exists(folder_path):
            return

        self.update_status("Scanning folder...")
        file_count = 0

        try:
            def scan_directory(path):
                nonlocal file_count
                try:
                    with os.scandir(path) as entries:
                        for entry in entries:
                            if entry.is_file() and entry.name.endswith('.docx') and not entry.name.startswith('~'):
                                self.loaded_files.append(entry.path)
                                file_count += 1
                                
                                if file_count % 50 == 0:
                                    self.update_status(f"Found {file_count} files...")
                                    self.root.update_idletasks()
                            elif entry.is_dir():
                                scan_directory(entry.path)
                except PermissionError:
                    self.log(f"⚠️ Permission denied accessing: {path}")
            
            scan_directory(folder_path)
            
            folder_stats = self.get_folder_stats()
            self.loaded_files_status.set(folder_stats)
            self.log(f"📁 {folder_stats} from folder: ")
            self.add_clickable_link(folder_path)
            self.update_status(f"Loaded {len(self.loaded_files)} files from folder")
        except Exception as e:
            self.log(f"❌ Error scanning folder: {str(e)}")
            self.update_status("Error scanning folder", True)

    def get_folder_stats(self):
        """Get total size and file count of loaded files"""
        total_size = 0
        file_count = len(self.loaded_files)
        
        for file_path in self.loaded_files:
            try:
                size = os.path.getsize(file_path)
                total_size += size
            except OSError:
                continue
        
        size_mb = round(total_size / 1024 / 1024, 1)
        return f"Loaded {file_count} files ({size_mb} MB total)"

    def load_files_from_zip(self, zip_path: str):
        """Load files from DocXScan ZIP archive with enhanced error handling"""
        self.loaded_files = []

        try:
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                temp_dir = tempfile.mkdtemp(prefix='docx_replace_')
                self.temp_directories.append(temp_dir)

                excel_files = [f for f in zip_ref.namelist() if f.endswith('.xlsx') and 'metadata' in f.lower()]

                if excel_files:
                    excel_filename = excel_files[0]
                    excel_data = zip_ref.read(excel_filename)

                    temp_excel = tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False)
                    temp_excel.write(excel_data)
                    temp_excel.close()

                    try:
                        df = pd.read_excel(temp_excel.name)
                        if 'File Path' in df.columns:
                            for zip_item in zip_ref.namelist():
                                if 'Matched_Files/' in zip_item and zip_item.endswith('.docx') and not zip_item.endswith('/'):
                                    file_data = zip_ref.read(zip_item)
                                    filename = os.path.basename(zip_item)
                                    extracted_path = os.path.join(temp_dir, filename)

                                    with open(extracted_path, 'wb') as f:
                                        f.write(file_data)

                                    self.loaded_files.append(extracted_path)

                            if not self.loaded_files:
                                for zip_item in zip_ref.namelist():
                                    if zip_item.endswith('.docx') and not zip_item.startswith('~') and not zip_item.endswith('/'):
                                        file_data = zip_ref.read(zip_item)
                                        filename = os.path.basename(zip_item)
                                        extracted_path = os.path.join(temp_dir, filename)

                                        with open(extracted_path, 'wb') as f:
                                            f.write(file_data)

                                        self.loaded_files.append(extracted_path)

                            folder_stats = self.get_folder_stats()
                            self.loaded_files_status.set(folder_stats + " from ZIP archive")
                            self.log(f"📦 Extracted {len(self.loaded_files)} files from ZIP: ")
                            self.add_clickable_link(zip_path)
                            self.log(f"📁 Temporary folder: ")
                            self.add_clickable_link(temp_dir)

                            self.selected_files_folder.set(temp_dir)
                            self.update_status(f"Loaded {len(self.loaded_files)} files from ZIP")

                        else:
                            messagebox.showerror("Error", "Excel file in ZIP must contain 'File Path' column")

                    finally:
                        os.unlink(temp_excel.name)
                else:
                    for zip_item in zip_ref.namelist():
                        if zip_item.endswith('.docx') and not zip_item.startswith('~') and not zip_item.endswith('/'):
                            file_data = zip_ref.read(zip_item)
                            filename = os.path.basename(zip_item)
                            extracted_path = os.path.join(temp_dir, filename)

                            with open(extracted_path, 'wb') as f:
                                f.write(file_data)

                            self.loaded_files.append(extracted_path)

                    folder_stats = self.get_folder_stats()
                    self.loaded_files_status.set(folder_stats + " from ZIP (no metadata)")
                    self.log(f"📦 Extracted {len(self.loaded_files)} files from ZIP without metadata: ")
                    self.add_clickable_link(zip_path)

                    self.selected_files_folder.set(temp_dir)
                    self.update_status(f"Loaded {len(self.loaded_files)} files from ZIP")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to load ZIP file:\n{str(e)}")
            self.log(f"❌ Error loading ZIP: {str(e)}")
            self.update_status(f"Error loading ZIP file", True)

    def cleanup_temp_files(self):
        """Clean up temporary extracted files with user confirmation"""
        if not self.temp_directories:
            self.log("ℹ️ No temporary directories to clean up")
            self.update_status("No temp files to clean")
            return

        if messagebox.askyesno("Clean Temp Files", 
                              f"Clean up {len(self.temp_directories)} temporary directories?\n\n"
                              "This will remove extracted files and clear loaded file list."):
            cleaned_count = 0
            for temp_dir in self.temp_directories:
                try:
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
                        cleaned_count += 1
                        self.log(f"🧹 Cleaned up temporary folder: {temp_dir}")
                except Exception as e:
                    self.log(f"⚠️ Could not clean up {temp_dir}: {str(e)}")

            self.temp_directories = []
            if cleaned_count > 0:
                self.log(f"✅ Cleaned up {cleaned_count} temporary directories")
                self.loaded_files = []
                self.loaded_files_status.set("No files loaded.")
                self.update_status(f"Cleaned {cleaned_count} temp directories")

    def load_files_from_excel(self, excel_path: str):
        """Load files from Excel with validation"""
        try:
            df = pd.read_excel(excel_path)
            if 'File Path' in df.columns:
                all_paths = df['File Path'].tolist()
                self.loaded_files = [path for path in all_paths if os.path.exists(path)]
                missing_count = len(all_paths) - len(self.loaded_files)

                folder_stats = self.get_folder_stats()
                self.loaded_files_status.set(folder_stats + " from Excel")
                self.log(f"📊 Loaded {len(self.loaded_files)} files from Excel scan results: ")
                self.add_clickable_link(excel_path)

                if missing_count > 0:
                    self.log(f"⚠️ {missing_count} files from Excel list were not found")

                self.update_status(f"Loaded {len(self.loaded_files)} files from Excel")
            else:
                messagebox.showerror("Error", "Excel file must contain 'File Path' column")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load Excel file:\n{str(e)}")
            self.update_status("Error loading Excel file", True)

    def load_replacement_file(self):
        """Load replacement file with enhanced validation"""
        path = filedialog.askopenfilename(
            title="Select Replacement JSON File", 
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
        )
        if path:
            self.load_replacement_file_direct(path)

    def load_replacement_file_direct(self, path: str):
        """Load replacement file directly from path"""
        try:
            with open(path, "r", encoding='utf-8') as f:
                self.replacement_map = json.load(f)

            # Validate replacement map
            errors = self.document_processor.validate_replacement_map(self.replacement_map)
            if errors:
                messagebox.showwarning("Validation Warnings", 
                                     f"Found {len(errors)} issues:\n" + "\n".join(errors[:5]))

            self.replacement_file_path.set(path)
            self.loaded_replacement_file.set(f"Loaded: {os.path.basename(path)} ✅ ({len(self.replacement_map)} replacements)")
            self.save_configuration()
            self.log(f"✅ Loaded {len(self.replacement_map)} replacements from: ")
            self.add_clickable_link(path)
            self.add_to_recent_replacements(path)
            
            # Validate regex patterns if regex mode is enabled
            if self.regex_mode.get():
                self.validate_regex_patterns()

        except Exception as e:
            messagebox.showerror("Error", f"Failed to load replacement file:\n{e}")
            self.loaded_replacement_file.set("❌ Failed to load replacement file.")
            self.update_status("Error loading replacement file", True)

    def add_to_recent_replacements(self, file_path: str):
        """Add replacement file to recent list"""
        if file_path in self.recent_replacements:
            self.recent_replacements.remove(file_path)
        self.recent_replacements.insert(0, file_path)
        self.recent_replacements = self.recent_replacements[:5]

    def on_regex_mode_changed(self):
        """Handle regex mode change"""
        if self.regex_mode.get() and self.replacement_map:
            self.validate_regex_patterns()

    def validate_patterns(self):
        """Validate all loaded patterns"""
        if not self.replacement_map:
            messagebox.showinfo("No Patterns", "No replacement patterns loaded.")
            return

        errors = self.document_processor.validate_replacement_map(self.replacement_map)

        if self.regex_mode.get():
            regex_errors = []
            for pattern in self.replacement_map.keys():
                try:
                    re.compile(pattern)
                except re.error as e:
                    regex_errors.append(f"'{pattern}': {e}")
            errors.extend(regex_errors)

        if errors:
            self.show_validation_errors(errors)
        else:
            messagebox.showinfo("Validation Complete", "All patterns are valid!")

    def show_validation_errors(self, errors):
        """Show validation errors in a window"""
        error_window = tk.Toplevel(self.root)
        error_window.title("Pattern Validation Errors")
        error_window.geometry("600x400")
        error_window.configure(bg=ModernColors.BG_PRIMARY)

        frame = tk.Frame(error_window, bg=ModernColors.BG_PRIMARY)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        text_widget = tk.Text(frame, wrap=tk.WORD, bg=ModernColors.CARD_BG, 
                             fg=ModernColors.TEXT_PRIMARY)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)

        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        error_text = "\n".join([f"• {error}" for error in errors])
        text_widget.insert(tk.END, f"Found {len(errors)} validation errors:\n\n{error_text}")
        text_widget.configure(state=tk.DISABLED)

    def validate_regex_patterns(self):
        """Validate regex patterns in replacement map"""
        invalid_patterns = []
        legal_token_warnings = []

        for old_pattern, new_pattern in self.replacement_map.items():
            try:
                re.compile(old_pattern)

                if self.regex_mode.get():
                    unescaped_chars = []
                    if '<' in old_pattern and '\\<' not in old_pattern:
                        unescaped_chars.append('<')
                    if '>' in old_pattern and '\\>' not in old_pattern:
                        unescaped_chars.append('>')
                    if '[' in old_pattern and '\\[' not in old_pattern:
                        unescaped_chars.append('[')
                    if ']' in old_pattern and '\\]' not in old_pattern:
                        unescaped_chars.append(']')
                    if '.' in old_pattern and '\\.' not in old_pattern:
                        unescaped_chars.append('.')

                    if unescaped_chars:
                        legal_token_warnings.append(f"Pattern '{old_pattern}' may need escaping for: {', '.join(unescaped_chars)}")

            except re.error:
                invalid_patterns.append(old_pattern)

        if invalid_patterns:
            messagebox.showerror(
                "Invalid Regex Patterns", 
                f"The following patterns are invalid:\n" + "\n".join(invalid_patterns[:5]) + 
                (f"\n... and {len(invalid_patterns)-5} more" if len(invalid_patterns) > 5 else "")
            )

        if legal_token_warnings:
            messagebox.showwarning(
                "Legal Token Pattern Warnings",
                "Potential issues with legal token patterns:\n" + 
                "\n".join(legal_token_warnings[:3]) +
                (f"\n... and {len(legal_token_warnings)-3} more" if len(legal_token_warnings) > 3 else "") +
                "\n\nConsider escaping special characters with backslashes."
            )

    def add_clickable_link(self, file_or_folder_path: str):
        """Add a clickable link to the console"""
        def open_item():
            try:
                normalized_path = os.path.normpath(file_or_folder_path)

                if not os.path.exists(normalized_path):
                    self.log(f"❌ Path does not exist: {normalized_path}")
                    return

                if platform.system() == "Windows":
                    os.startfile(normalized_path)
                elif platform.system() == "Darwin":  # macOS
                    subprocess.run(["open", normalized_path])
                else:  # Linux
                    subprocess.run(["xdg-open", normalized_path])

                item_type = "folder" if os.path.isdir(normalized_path) else "file"
                self.log(f"📂 Opened {item_type}: {os.path.basename(normalized_path)}")

            except Exception as e:
                self.log(f"❌ Could not open item: {str(e)}")
                try:
                    parent_dir = os.path.dirname(normalized_path)
                    if os.path.exists(parent_dir):
                        os.startfile(parent_dir)
                        self.log(f"📂 Opened parent folder instead: {parent_dir}")
                except Exception:
                    self.log(f"❌ Could not open parent folder either")

        start_pos = self.console.text_widget.index(tk.END + "-1c")
        display_text = os.path.basename(file_or_folder_path) if os.path.isfile(file_or_folder_path) else file_or_folder_path
        self.console.text_widget.insert(tk.END, display_text + "\n")
        end_pos = self.console.text_widget.index(tk.END + "-1c")

        link_tag = f"link_{datetime.now().strftime('%H%M%S%f')}"
        self.console.text_widget.tag_add(link_tag, start_pos, end_pos)

        self.console.text_widget.tag_config(link_tag, 
                                           foreground=ModernColors.PRIMARY, 
                                           underline=True)

        self.console.text_widget.tag_bind(link_tag, "<Button-1>", lambda e: open_item())
        self.console.text_widget.tag_bind(link_tag, "<Enter>", lambda e: self.console.text_widget.configure(cursor="hand2"))
        self.console.text_widget.tag_bind(link_tag, "<Leave>", lambda e: self.console.text_widget.configure(cursor=""))

        self.console.text_widget.see(tk.END)

    def create_replacement_template(self):
        """Create a template replacement JSON file"""
        template = {
            "<<FileService.": "<<NewFileService.",
            "</ff>": "<<PAGE_BREAK>>",
            "</pp>": "<<HARD_RETURN>>",
            "<backspace>": "<<BACKSPACE>>",
            "<<STNDRDTH": "<<STANDARD_TH",
            "<c>": "<<CENTER>>",
            "<u>": "<<UNDERLINE>>",
            "<i>": "<<ITALIC>>",
            "<bold>": "<<BOLD>>",
            "</nobullet>": "<<NO_BULLET>>",
            "[[MCOMPUTEINTO(<<": "<<MCOMPUTE_INTO(",
            "[[SCOMPUTEINTO(": "<<SCOMPUTE_INTO(",
            "[[ABORTIIF": "<<ABORT_IF",
            "PROMTINTO(": "<<PROMPT_INTO(",
            "PROMTINTOIIF(": "<<PROMPT_INTO_IF(",
            "<<Checklist.": "<<CHECKLIST.",
            "TABLE(": "<<TABLE(",
            "<<jfig": "<<JFIG",
            "{ATTY": "<<ESIGN_ATTORNEY"
        }

        if self.regex_mode.get():
            template = {
                r"<<FileService\.\\w*": "<<NewFileService.{{match}}",
                r"<<BLTO\\d+": "<<BULLET_ORDERED_{{match}}>>",
                r"<<BLT#\\d+": "<<BULLET_NUMBERED_{{match}}>>",
                r"\\[\\[(\\w+)COMPUTEINTO\\(": "<<{{match}}_INTO(",
                r"PROMT(\\w*)\\(": "<<PROMPT_{{match}}(",
                r"<<Special\\.(\\w*)": "<<SPECIAL.{{match}}",
                r"<<Tracker\\.(\\w+)>>~(\\w+):": "<<TRACKER.{{match}}_FORMAT>>",
                r"\\{ATTY(\\w*)": "<<ESIGN_{{match}}",
                r"<(\\w+)>": "<<{{match}}>>",
                r"</(\\w+)>": "<<END_{{match}}>>",
                r"[+-]\\d+\\|<<Special\\.ToDay:": "<<SPECIAL.DATE_OFFSET_{{match}}>>"
            }

        path = filedialog.asksaveasfilename(
            title="Save Replacement Template",
            defaultextension=".json",
            filetypes=[("JSON Files", "*.json")]
        )

        if path:
            try:
                with open(path, "w", encoding='utf-8') as f:
                    json.dump(template, f, indent=2)
                self.log(f"📄 Created template: {os.path.basename(path)}")
                messagebox.showinfo("Success", f"Template created successfully!\n{path}")
                self.update_status("Template created successfully")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to create template:\n{str(e)}")

    def create_output_copy(self, file_path: str, output_root: str, session_timestamp: str = None) -> Tuple[str, str]:
        """Create a copy of the file in output directory for modification"""
        if session_timestamp is None:
            session_timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        output_dir = os.path.join(output_root, f"modified_{session_timestamp}")

        try:
            os.makedirs(output_dir, exist_ok=True)
        except PermissionError:
            user_docs = os.path.expanduser("~/Documents")
            output_dir = os.path.join(user_docs, "DocReplace_Output", f"modified_{session_timestamp}")
            os.makedirs(output_dir, exist_ok=True)
            self.log(f"⚠️ Using fallback output location: {output_dir}")

        output_filename = os.path.basename(file_path)
        base_name, ext = os.path.splitext(output_filename)
        counter = 1
        output_path = os.path.join(output_dir, output_filename)

        while os.path.exists(output_path):
            output_filename = f"{base_name}_{counter}{ext}"
            output_path = os.path.join(output_dir, output_filename)
            counter += 1

        try:
            shutil.copy2(file_path, output_path)
            self.log(f"📋 Created working copy: {os.path.basename(file_path)} → {output_filename}")
        except Exception as e:
            self.log(f"❌ Copy failed for {os.path.basename(file_path)}: {str(e)}")
            raise

        return output_dir, output_path

    def run_replacement_threaded(self):
        """Start replacement process in a separate thread"""
        self.stop_processing = False
        self.eta_text.set("⏱️ Calculating...")
        self.progress_bar.set_progress(0, "Starting process...")
        threading.Thread(target=self.process_replacements, daemon=True).start()

    def process_replacements(self):
        """Main replacement processing function with enhanced error handling and progress tracking"""
        self.stop_processing = False

        # Enable stop button, disable start button during processing
        self.root.after(0, lambda: self.stop_button.configure(state="normal"))
        self.root.after(0, lambda: self.start_button.configure(state="disabled"))

        try:
            # Validation
            if not self.loaded_files:
                self.root.after(0, lambda: messagebox.showerror("No Files", "Please load files to process first."))
                return

            if not self.replacement_map:
                self.root.after(0, lambda: messagebox.showerror("No Replacements", "Please load a replacement file first."))
                return

            mode = self.replacement_mode.get()
            backup_root = self.backup_folder.get()

            if mode != "Dry Run (preview only)" and "Modified Copies" in mode and not backup_root:
                self.root.after(0, lambda: messagebox.showerror("No Output Folder", "Please select an output folder."))
                return

            # Initialize progress
            total_files = len(self.loaded_files)
            self.root.after(0, lambda: self.progress_bar.progress_bar.configure(maximum=100, value=0))
            self.root.after(0, lambda: self.log(f"🚀 Starting {mode} on {total_files} files..."))
            self.root.after(0, lambda: self.update_status(f"Processing {total_files} files..."))

            processed_files = 0
            modified_files = 0
            total_replacements = 0
            current_output_dir = None
            session_timestamp = None
            start_time = time.time()

            for i, file_path in enumerate(self.loaded_files):
                # Check for stop request
                if self.stop_processing:
                    self.root.after(0, lambda: self.log("🛑 Processing stopped by user"))
                    break

                if not os.path.exists(file_path):
                    self.root.after(0, lambda f=file_path: self.log(f"⚠️ File not found: {f}"))
                    continue

                try:
                    # Update progress with ETA
                    elapsed = time.time() - start_time
                    eta_seconds = 0
                    
                    if i > 0:
                        avg_time = elapsed / i
                        remaining_files = total_files - i
                        eta_seconds = remaining_files * avg_time
                    
                    current_progress = i + 1
                    self.root.after(0, lambda: self.update_progress(current_progress, total_files, eta_seconds, file_path))

                    # Load document
                    doc = Document(file_path)
                    
                    try:
                        # Perform replacements
                        replacements_made, replacement_details = self.document_processor.perform_replacement_in_doc(
                            doc, file_path, self.replacement_map, self.regex_mode.get())
                        
                        if replacements_made > 0:
                            if mode == "Dry Run (preview only)":
                                self.root.after(0, lambda f=file_path, r=replacements_made: 
                                               self.log(f"🔍 Would modify {os.path.basename(f)}: {r} replacements"))
                                modified_files += 1
                            else:
                                if "Modified Copies" in mode:
                                    # Create output copy (originals stay untouched)
                                    if current_output_dir is None:
                                        session_timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                        current_output_dir, output_file_path = self.create_output_copy(
                                            file_path, backup_root, session_timestamp)
                                    else:
                                        _, output_file_path = self.create_output_copy(
                                            file_path, backup_root, session_timestamp)

                                    # Save modified document to the OUTPUT copy (not original)
                                    doc.save(output_file_path)
                                    modified_files += 1
                                    total_replacements += replacements_made

                                    self.root.after(0, lambda f=file_path, r=replacements_made: 
                                                   self.log(f"✅ Created modified copy of {os.path.basename(f)}: {r} replacements"))
                                else:
                                    # In-place replacement mode
                                    doc.save(file_path)
                                    modified_files += 1
                                    total_replacements += replacements_made

                                    self.root.after(0, lambda f=file_path, r=replacements_made: 
                                                   self.log(f"✅ Modified {os.path.basename(f)}: {r} replacements"))
                        else:
                            self.root.after(0, lambda f=file_path: self.log(f"➖ No changes needed: {os.path.basename(f)}"))

                        processed_files += 1

                    finally:
                        # Always clean up document from memory
                        del doc

                except Exception as e:
                    error_msg = str(e)
                    self.root.after(0, lambda f=file_path, error=error_msg: 
                                   self.log(f"❌ Error processing {os.path.basename(f)}: {error}"))
                    self.logger.error(f"Error processing {file_path}: {error_msg}")
                    self.logger.error(traceback.format_exc())

            # Add to backup history if outputs were created
            if current_output_dir and mode != "Dry Run (preview only)" and not self.stop_processing:
                self.backup_history.append({
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'backup_dir': current_output_dir,
                    'files_modified': modified_files,
                    'total_replacements': total_replacements,
                    'mode': mode
                })
                self.root.after(0, lambda: self.undo_button.configure(state="normal"))

            # Final summary
            elapsed_total = time.time() - start_time
            if self.stop_processing:
                self.root.after(0, lambda: self.log(f"\n🛑 Processing Stopped:"))
                self.root.after(0, lambda: self.log(f"   • Files processed: {processed_files}"))
                self.root.after(0, lambda: self.log(f"   • Files modified: {modified_files}"))
                self.root.after(0, lambda: self.log(f"   • Total replacements: {total_replacements}"))
                self.root.after(0, lambda: self.log(f"   • Time elapsed: {elapsed_total:.1f}s"))
                if current_output_dir:
                    self.root.after(0, lambda: self.log(f"   • Output folder created: "))
                    self.root.after(0, lambda: self.add_clickable_link(current_output_dir))
                self.root.after(0, lambda: self.update_status("Processing stopped by user"))
                self.root.after(0, lambda: self.progress_bar.set_progress(0, "Stopped"))
            elif mode == "Dry Run (preview only)":
                self.root.after(0, lambda: self.log(f"\n📋 Dry Run Complete:"))
                self.root.after(0, lambda: self.log(f"   • Files processed: {processed_files}"))
                self.root.after(0, lambda: self.log(f"   • Files that would be modified: {modified_files}"))
                self.root.after(0, lambda: self.log(f"   • Time elapsed: {elapsed_total:.1f}s"))
                self.root.after(0, lambda: self.update_status("Dry run completed"))
                self.root.after(0, lambda: self.progress_bar.set_progress(100, "Dry run complete"))
            else:
                self.root.after(0, lambda: self.log(f"\n🎉 Replacement Complete:"))
                self.root.after(0, lambda: self.log(f"   • Files processed: {processed_files}"))
                self.root.after(0, lambda: self.log(f"   • Files modified: {modified_files}"))
                self.root.after(0, lambda: self.log(f"   • Total replacements: {total_replacements}"))
                self.root.after(0, lambda: self.log(f"   • Time elapsed: {elapsed_total:.1f}s"))
                if current_output_dir:
                    self.root.after(0, lambda: self.log(f"   • Output folder created: "))
                    self.root.after(0, lambda: self.add_clickable_link(current_output_dir))
                self.root.after(0, lambda: self.update_status(f"Completed: {modified_files} files modified"))
                self.root.after(0, lambda: self.progress_bar.set_progress(100, f"Complete: {modified_files} files modified"))
                self.root.after(0, lambda: self.eta_text.set(""))  # Clear ETA when complete

            # Log operation to file
            self.logger.info(f"Operation completed - Mode: {mode}, Files: {processed_files}, "
                           f"Modified: {modified_files}, Replacements: {total_replacements}, "
                           f"Time: {elapsed_total:.1f}s")

            self.save_configuration()

        except Exception as e:
            self.root.after(0, lambda: self.log(f"❌ Critical error: {str(e)}"))
            self.root.after(0, lambda: messagebox.showerror("Critical Error", f"Processing failed:\n{str(e)}"))
            self.logger.error(f"Critical error in process_replacements: {str(e)}")
            self.logger.error(traceback.format_exc())

        finally:
            # Clear ETA display when done
            self.root.after(0, lambda: self.eta_text.set(""))
            # Always re-enable buttons when done
            self.root.after(0, lambda: self.stop_button.configure(state="disabled"))
            self.root.after(0, lambda: self.start_button.configure(state="normal"))
            
            # Clean up memory after large operations
            import gc
            gc.collect()

    def save_batch_config(self):
        """Save current settings as a batch configuration"""
        config = {
            'replacement_file': self.replacement_file_path.get(),
            'output_folder': self.backup_folder.get(),
            'mode': self.replacement_mode.get(),
            'regex_enabled': self.regex_mode.get(),
            'auto_backup': self.auto_backup.get(),
            'timestamp': datetime.now().isoformat(),
            'version': '3.0'
        }

        path = filedialog.asksaveasfilename(
            title="Save Batch Configuration",
            defaultextension=".json",
            filetypes=[("JSON Files", "*.json")]
        )

        if path:
            try:
                with open(path, 'w') as f:
                    json.dump(config, f, indent=2)
                self.log(f"💾 Saved batch configuration: {os.path.basename(path)}")
                self.update_status("Batch configuration saved")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save configuration:\n{e}")

    def load_batch_config(self):
        """Load batch configuration"""
        path = filedialog.askopenfilename(
            title="Load Batch Configuration",
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
        )

        if path:
            try:
                with open(path, 'r') as f:
                    config = json.load(f)

                # Apply configuration
                if 'replacement_file' in config and os.path.exists(config['replacement_file']):
                    self.replacement_file_path.set(config['replacement_file'])
                    self.load_replacement_file_direct(config['replacement_file'])

                if 'output_folder' in config:
                    self.backup_folder.set(config['output_folder'])

                if 'mode' in config:
                    self.replacement_mode.set(config['mode'])

                if 'regex_enabled' in config:
                    self.regex_mode.set(config['regex_enabled'])

                if 'auto_backup' in config:
                    self.auto_backup.set(config['auto_backup'])

                self.log(f"📂 Loaded batch configuration: {os.path.basename(path)}")
                self.update_status("Batch configuration loaded")

            except Exception as e:
                messagebox.showerror("Error", f"Failed to load configuration:\n{e}")

    def show_settings(self):
        """Show settings dialog"""
        settings_window = tk.Toplevel(self.root)
        settings_window.title("Settings")
        settings_window.geometry("400x300")
        settings_window.configure(bg=ModernColors.BG_PRIMARY)
        settings_window.transient(self.root)
        settings_window.grab_set()

        # Auto-backup setting
        auto_backup_frame = tk.Frame(settings_window, bg=ModernColors.BG_PRIMARY)
        auto_backup_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Checkbutton(auto_backup_frame, text="Auto-backup enabled", 
                       variable=self.auto_backup,
                       font=(self.font_family, 10),
                       bg=ModernColors.BG_PRIMARY,
                       fg=ModernColors.TEXT_PRIMARY,
                       selectcolor=ModernColors.BG_SECONDARY,
                       activebackground=ModernColors.BG_PRIMARY,
                       activeforeground=ModernColors.TEXT_PRIMARY).pack(anchor=tk.W)

        # Default output folder
        tk.Label(settings_window, text="Default Output Folder:",
                font=(self.font_family, 10, "bold"),
                bg=ModernColors.BG_PRIMARY,
                fg=ModernColors.TEXT_PRIMARY).pack(anchor=tk.W, padx=10, pady=(10, 0))

        folder_frame = tk.Frame(settings_window, bg=ModernColors.BG_PRIMARY)
        folder_frame.pack(fill=tk.X, padx=10, pady=5)

        default_folder = tk.StringVar(value=self.settings.get_setting("default_output_folder"))
        folder_entry = tk.Entry(folder_frame, textvariable=default_folder, 
                               bg=ModernColors.BG_SECONDARY, fg=ModernColors.TEXT_PRIMARY,
                               state="readonly")
        folder_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))
        
        browse_btn = ModernButton(folder_frame, "Browse", 
                                 lambda: self.browse_default_folder(default_folder), style="secondary")
        browse_btn.pack(side=tk.RIGHT)

        # Buttons
        button_frame = tk.Frame(settings_window, bg=ModernColors.BG_PRIMARY)
        button_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=10)

        save_btn = ModernButton(button_frame, "Save", 
                               lambda: self.save_settings(settings_window, default_folder))
        save_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        cancel_btn = ModernButton(button_frame, "Cancel", settings_window.destroy, style="secondary")
        cancel_btn.pack(side=tk.RIGHT)

    def browse_default_folder(self, folder_var):
        """Browse for default folder"""
        path = filedialog.askdirectory(title="Select Default Output Folder")
        if path:
            folder_var.set(path)

    def save_settings(self, window, default_folder):
        """Save settings"""
        settings = {
            "auto_backup": str(self.auto_backup.get()),
            "default_output_folder": default_folder.get()
        }
        self.settings.save_config(settings)
        window.destroy()
        self.update_status("Settings saved")

    def view_logs(self):
        """Open log viewer"""
        log_window = tk.Toplevel(self.root)
        log_window.title("Log Viewer")
        log_window.geometry("800x600")
        log_window.configure(bg=ModernColors.BG_PRIMARY)

        # Create text widget with scrollbar
        frame = tk.Frame(log_window, bg=ModernColors.BG_PRIMARY)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        text_widget = tk.Text(frame, wrap=tk.WORD, bg=ModernColors.CARD_BG, 
                             fg=ModernColors.TEXT_PRIMARY, font=("Consolas", 9))
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)

        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Load log content
        try:
            log_file = f'logs/docx_replace_{datetime.now().strftime("%Y%m%d")}.log'
            if os.path.exists(log_file):
                with open(log_file, 'r') as f:
                    text_widget.insert(tk.END, f.read())
            else:
                text_widget.insert(tk.END, "No log file found for today.")
        except Exception as e:
            text_widget.insert(tk.END, f"Error loading log file: {e}")

        text_widget.configure(state=tk.DISABLED)

    def show_about(self):
        """Show about dialog"""
        about_text = """DocXReplace v3.0 Professional

© 2025 Hrishik Kunduru. All rights reserved.

A professional tool for automated document replacement with 
legal token migration support.

Features:
• Ultra-modern interface matching Document Tools Suite
• Batch document processing
• Regular expression support
• Safe backup system
• Professional logging
• Clickable file links
• Keyboard shortcuts
• Comprehensive help system

For support and updates, contact the developer."""

        messagebox.showinfo("About DocXReplace v3.0", about_text)

    def show_token_help(self):
        """Show general token help dialog"""
        help_text = """
GENERAL LEGAL DOCUMENT TOKENS

Text Formatting Tokens:
• <<FileService. → Fileservice
• <c> → Center
• <u> → Underline  
• <i> → Italic
• <bold> → Bold
• <pcase> → Proper case
• <lcase> → Lower case
• <ucase> → Upper case
• <fontsize → Font Size
• <nobullet> → No Bullet

Document Structure Tokens:
• </ff> → Page Break
• </pp> → Hard Return
• <backspace> → Backspace
• <s1> → Section 1
• <s2> → Section 2

Numbering and Bullets:
• <<BLTO1 → Adds numbering to the text in the doc
• <<BLT#1 → Same but with #1
• <<BLTO2>>/<<BLT#2>> → Resets the numbering on a new doc
• <<BLT#2 → Same but with #2
• <<BLTA4 → Adds lettering

Computation Tokens:
• [[MCOMPUTEINTO(<< → MCOMPUTE INTO
• [[SCOMPUTEINTO( → SCOMPUTE INTO
• [[ABORTIIF → Abort if

Prompt Tokens:
• PROMTINTO( → PROMTINTO
• PROMTINTOIIF( → PROMTINTOIIF
• PROMTINTOLIST( → PROMTINTOLIST
• PROMTINTOIIFLIST( → PROMTINTOIIFLIST
• PROMTFORM( → PROMTFORM

Special Function Tokens:
• <<Checklist. → CHECKLIST
• TABLE( → TABLE
• <<jfig → JFIG
• jfig → JFIG_General
• {ATTY → ESIGN
• <<Special. → SPECIAL
• <<STNDRDTH → STNDRD Add "TH"

Date and Tracker Tokens:
• +91|<<Special.ToDay: → Add 91 days to today
• -91|<<Special.ToDay: → Subtract 91 days from today
• +2|<<Special.ToDay: → Add 2 days to today
• -2|<<Special.ToDay: → Subtract 2 days from today
• <<Tracker.MortDate>>~MMMM dd: → Mortgage date in Month DD format
• <<Tracker.MortDate>>~MM-dd-yyyy: → Mortgage date in MM-DD-YYYY format
• <<Tracker.MortDate>>~ddd: → Mortgage date as day of week
• <<Tracker.OriginalPrincipal>>~##: → Original principal amount

Token Usage Tips:
• Tokens are case-sensitive
• Some tokens require specific syntax (like closing brackets)
• Date tokens use + or - followed by number for date arithmetic
• Tracker tokens use ~format after >> to specify output format
• Test tokens with Dry Run before applying to documents
• Use exact spelling and punctuation as shown above
"""

        self.show_help_window("Legal Token Reference Guide", help_text)

    def show_regex_help(self):
        """Show regex help dialog"""
        help_text = """
REGEX PATTERNS FOR LEGAL DOCUMENT TOKEN MIGRATION

Common Legal Token Patterns:
• \\< and \\> - Escape angle brackets for literal matching
• \\. - Escape periods in tokens like <<FileService.
• \\[ and \\] - Escape square brackets for [[COMPUTEINTO
• \\w+ - One or more word characters
• \\d+ - One or more digits
• [^)]* - Any character except closing parenthesis

Legal Document Token Examples:
• <<FileService\\.\\w* → <<NewFileService.{{match}}
  Matches: <<FileService.Save, <<FileService.Load

• <<BLT(#|O)(\\d+) → <<BULLET_{{match}}_{{match}}>>
  Matches: <<BLT#1, <<BLTO2 → <<BULLET_#_1>>, <<BULLET_O_2>>

• \\[\\[(\\w+)COMPUTEINTO\\( → <<{{match}}_COMPUTE_INTO(
  Matches: [[MCOMPUTEINTO(, [[SCOMPUTEINTO(

• <(\\w+)> → <<{{match}}>>
  Matches: <bold>, <italic> → <<bold>>, <<italic>>

• PROMT(\\w*)INTO(\\w*)\\( → <<PROMPT_{{match}}_INTO_{{match}}(
  Matches: PROMTINTO(, PROMTINTOIIF( with variable parts

Special Date/Tracker Patterns:
• ([+-]\\d+)\\|<<Special\\.ToDay: → <<SPECIAL.DATE_OFFSET_{{match}}>>
  Matches: +91|<<Special.ToDay:, -2|<<Special.ToDay:

• <<Tracker\\.(\\w+)>>~(\\w+): → <<TRACKER.{{match}}_FORMAT>>
  Matches: <<Tracker.MortDate>>~MMMM dd:

Replacement Syntax:
• {{match}} - Replaced with captured group or full match
• Use multiple {{match}} for multiple capture groups
• Groups are replaced in order: first {{match}} = group 1, etc.

Legal Token Migration Tips:
• Always escape special characters: < > [ ] . ( ) 
• Test with simple tokens first (like <bold> → <<BOLD>>)
• Use capture groups for variable parts of tokens
• Preview with Dry Run before applying to documents
• Consider token context (some may appear in different formats)
"""

        self.show_help_window("Legal Token Regex Help", help_text)

    def show_help_window(self, title: str, content: str):
        """Generic help window creator"""
        help_window = tk.Toplevel(self.root)
        help_window.title(title)
        help_window.geometry("750x650")
        help_window.configure(bg=ModernColors.BG_PRIMARY)
        help_window.transient(self.root)

        # Add scrollbar
        frame = tk.Frame(help_window, bg=ModernColors.BG_PRIMARY)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        text_widget = tk.Text(frame, wrap=tk.WORD, bg=ModernColors.CARD_BG, 
                             fg=ModernColors.TEXT_PRIMARY, font=("Consolas", 10), padx=10, pady=10)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)

        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        text_widget.insert(tk.END, content)
        text_widget.configure(state=tk.DISABLED)

    def undo_last_operation(self):
        """Undo the last replacement operation"""
        if not self.backup_history:
            messagebox.showinfo("No Operations", "No operations to undo.")
            return

        last_backup = self.backup_history[-1]
        backup_dir = last_backup['backup_dir']

        if not os.path.exists(backup_dir):
            messagebox.showerror("Backup Not Found", f"Backup directory not found:\n{backup_dir}")
            return

        result = messagebox.askyesno(
            "Confirm Undo", 
            f"Undo last operation?\n"
            f"Timestamp: {last_backup['timestamp']}\n"
            f"Files modified: {last_backup['files_modified']}\n"
            f"Mode: {last_backup.get('mode', 'Unknown')}\n\n"
            f"This will restore files from backup."
        )

        if result:
            try:
                restored_count = 0
                for root_dir, _, files in os.walk(backup_dir):
                    for file in files:
                        if file.endswith('.docx'):
                            backup_file_path = os.path.join(root_dir, file)

                            original_path = None
                            for loaded_file in self.loaded_files:
                                if os.path.basename(loaded_file) == file:
                                    original_path = loaded_file
                                    break

                            if original_path and os.path.exists(original_path):
                                shutil.copy2(backup_file_path, original_path)
                                restored_count += 1

                self.backup_history.pop()
                self.log(f"↩️ Undo completed: {restored_count} files restored")
                self.update_status(f"Undo completed: {restored_count} files restored")

                if not self.backup_history:
                    self.undo_button.configure(state="disabled")

            except Exception as e:
                messagebox.showerror("Undo Failed", f"Failed to undo operation:\n{str(e)}")
                self.logger.error(f"Undo operation failed: {str(e)}")

    def show_backup_history(self):
        """Show backup history window with enhanced information"""
        if not self.backup_history:
            messagebox.showinfo("No History", "No backup history available.")
            return

        history_window = tk.Toplevel(self.root)
        history_window.title("Operation History")
        history_window.geometry("700x450")
        history_window.configure(bg=ModernColors.BG_PRIMARY)
        history_window.transient(self.root)

        # Create treeview for history
        columns = ("Timestamp", "Mode", "Files Modified", "Replacements", "Location")
        tree = ttk.Treeview(history_window, columns=columns, show="headings", height=15)

        # Configure columns
        tree.heading("Timestamp", text="Timestamp")
        tree.heading("Mode", text="Mode")
        tree.heading("Files Modified", text="Files")
        tree.heading("Replacements", text="Replacements")
        tree.heading("Location", text="Output Location")

        tree.column("Timestamp", width=140)
        tree.column("Mode", width=120)
        tree.column("Files Modified", width=80)
        tree.column("Replacements", width=100)
        tree.column("Location", width=200)

        # Populate with backup history
        for i, backup in enumerate(reversed(self.backup_history)):
            tree.insert("", "end", values=(
                backup['timestamp'],
                backup.get('mode', 'Unknown'),
                backup['files_modified'],
                backup['total_replacements'],
                backup['backup_dir']
            ))

        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Add buttons
        button_frame = tk.Frame(history_window, bg=ModernColors.BG_PRIMARY)
        button_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=10)

        def open_selected():
            selection = tree.selection()
            if not selection:
                messagebox.showwarning("No Selection", "Please select an operation to open its location.")
                return

            item = tree.item(selection[0])
            backup_path = item['values'][4]

            try:
                if platform.system() == "Windows":
                    os.startfile(backup_path)
                elif platform.system() == "Darwin":
                    subprocess.run(["open", backup_path])
                else:
                    subprocess.run(["xdg-open", backup_path])
            except Exception as e:
                messagebox.showerror("Error", f"Could not open location:\n{str(e)}")

        open_btn = ModernButton(button_frame, "Open Location", open_selected, style="secondary")
        open_btn.pack(side=tk.LEFT)
        
        close_btn = ModernButton(button_frame, "Close", history_window.destroy, style="secondary")
        close_btn.pack(side=tk.RIGHT)

    def load_configuration(self):
        """Load configuration from file"""
        config = self.settings.load_config()

        if config:
            self.selected_files_folder.set(config.get("files_folder", ""))
            self.backup_folder.set(config.get("backup_folder", ""))
            self.replacement_mode.set(config.get("replacement_mode", "Create Modified Copies (originals untouched)"))
            self.regex_mode.set(config.get("regex_mode", "False") == "True")

            replacement_file = config.get("replacement_file", "")
            if replacement_file and os.path.exists(replacement_file):
                self.load_replacement_file_direct(replacement_file)

        # Load recent files
        recent_files_str = self.settings.get_setting("recent_files", "")
        if recent_files_str:
            try:
                self.recent_files = json.loads(recent_files_str)
                self.update_recent_menu()
            except Exception:
                pass

    def save_configuration(self):
        """Save configuration to file"""
        config = {
            "files_folder": self.selected_files_folder.get(),
            "backup_folder": self.backup_folder.get(),
            "replacement_mode": self.replacement_mode.get(),
            "replacement_file": self.replacement_file_path.get(),
            "regex_mode": str(self.regex_mode.get()),
            "recent_files": json.dumps(self.recent_files)
        }

        self.settings.save_config(config)

    def on_closing(self):
        """Handle application closing"""
        try:
            # Save configuration
            self.save_configuration()

            # Clean up temp files
            self.cleanup_temp_files()

            # Log shutdown
            self.logger.info("DocXReplace v3.0 shutting down")

        except Exception as e:
            self.logger.error(f"Error during shutdown: {str(e)}")
        finally:
            self.root.destroy()

    def run(self):
        """Start the application"""
        # Set window protocol for closing
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        # Center window
        self.center_window()

        # Start the main loop
        self.log("🚀 DocXReplace v3.0 Professional ready!")
        self.log("📝 Load files and replacement patterns to begin processing")
        self.root.mainloop()

def main():
    """Main entry point"""
    try:
        app = DocXReplaceApp()
        app.run()
    except Exception as e:
        # Emergency error handling
        print(f"Critical error starting application: {e}")
        try:
            import tkinter.messagebox as mb
            mb.showerror("Critical Error", f"Failed to start DocXReplace:\n{e}")
        except Exception:
            pass

def run_docxreplace():
    """Entry point for the launcher"""
    app = DocXReplaceApp()
    app.run()

if __name__ == "__main__":
    main()


