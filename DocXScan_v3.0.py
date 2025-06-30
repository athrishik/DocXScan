#!/usr/bin/env python
# coding: utf-8

"""
DocXScan v3.0 - Ultra Modern Document Scanner
Copyright 2025 Hrishik Kunduru. All rights reserved.

Professional document scanner with intelligent token detection and modern UI.
"""

import subprocess
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import tempfile
import os
import re
import zipfile
import pandas as pd
from datetime import datetime
from docx import Document
import json
import configparser
import sys
import ctypes
import threading
import shutil
from pathlib import Path
import traceback
import platform
import logging
import time
from typing import Dict, List, Tuple, Optional, Any

# DPI awareness and console hiding for Windows
if platform.system() == "Windows":
    try:
        ctypes.windll.shcore.SetProcessDpiAwareness(2)
        ctypes.windll.user32.ShowWindow(ctypes.windll.kernel32.GetConsoleWindow(), 0)
    except:
        pass

class ModernColors:
    """Premium color palette matching Document Tools Suite exactly"""
    BG_PRIMARY = "#0A0E1A"      # Dark background
    BG_SECONDARY = "#0F1419"     # Darker variation
    
    GLASS_BG = "#1A1D29"        # Glass effect backgrounds
    GLASS_BORDER = "#2A2D3A"    # Glass borders
    GLASS_HOVER = "#1F2235"     # Glass hover state
    
    CARD_BG = "#161925"         # Card backgrounds
    CARD_BORDER = "#252837"     # Card borders
    CARD_HOVER = "#1C1F2E"      # Card hover state
    
    PRIMARY = "#2563EB"         # Blue theme
    PRIMARY_HOVER = "#1D4ED8"   # Darker blue for hover
    SUCCESS = "#10B981"         # Green theme
    SUCCESS_HOVER = "#059669"   # Darker green for hover
    WARNING = "#C56C86"         # Pink tone for warnings
    ERROR = "#FF7582"           # Coral red for errors
    
    TEXT_PRIMARY = "#FFFFFF"    # Pure white text 
    TEXT_SECONDARY = "#E5E7EB"  # Light gray
    TEXT_TERTIARY = "#C5A7A7"   # Warm gray with pink undertone 
    TEXT_MUTED = "#8B7D8B"      # Muted purple-gray 
    
    FOCUS_RING = "#2563EB"      # Blue focus ring
    
    # Gradient colors for decorative elements
    GRADIENT_1 = "#2563EB"      # Blue
    GRADIENT_2 = "#725A7A"      # Purple
    GRADIENT_3 = "#10B981"      # Green
    GRADIENT_4 = "#FF7582"      # Coral

class BlurredBackground(tk.Canvas):
    """Creates a gradient background effect"""
    
    def __init__(self, parent, width, height):
        super().__init__(parent, width=width, height=height, highlightthickness=0, bd=0)
        self.configure(bg=ModernColors.BG_PRIMARY)
        self.width = width
        self.height = height
        self.create_gradient_background()
    
    def create_gradient_background(self):
        """Create gradient background with darker tones"""
        for i in range(self.height):
            ratio = i / self.height
            # Create a smooth gradient from very dark to slightly lighter dark
            color_val = int(10 + (25 - 10) * (1 - ratio))
            color = f"#{color_val:02x}{color_val + 2:02x}{color_val + 8:02x}"
            self.create_line(0, i, self.width, i, fill=color, width=1)
        
        # Add decorative circles with sunset colors on dark background
        self.create_blur_circle(120, 80, 50, ModernColors.GRADIENT_1)      # Blue
        self.create_blur_circle(self.width - 100, 150, 40, ModernColors.GRADIENT_3)  # Pink
        self.create_blur_circle(250, self.height - 120, 60, ModernColors.GRADIENT_2)  # Purple

    def create_blur_circle(self, x, y, radius, color):
        """Create decorative circle with specified color"""
        for i in range(3):  # Create multiple layers for blur effect
            r = radius - i * 15
            if r > 0:
                alpha_color = color
                if i == 0:
                    alpha_color = color
                elif i == 1:
                    alpha_color = ModernColors.GRADIENT_2
                else:
                    alpha_color = ModernColors.GRADIENT_4
                    
                self.create_oval(x - r, y - r, x + r, y + r,
                               fill=alpha_color, outline="", stipple="gray25")

class ModernCard(tk.Frame):
    """Modern card component"""
    
    def __init__(self, parent, title, width=None, height=None):
        super().__init__(parent)
        self.title = title
        self.font_family = self.get_system_font()
        
        self.setup_card(width, height)
        self.create_header()
    
    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")
    
    def setup_card(self, width, height):
        """Setup card styling"""
        self.configure(
            bg=ModernColors.CARD_BG,
            relief="flat",
            bd=0,
            highlightthickness=1,
            highlightbackground=ModernColors.CARD_BORDER
        )
        
        if width and height:
            self.configure(width=width, height=height)
            self.pack_propagate(False)
    
    def create_header(self):
        """Create card header"""
        header_frame = tk.Frame(self, bg=ModernColors.CARD_BG)
        header_frame.pack(fill=tk.X, padx=16, pady=(12, 8))
        
        # Title
        tk.Label(header_frame, text=self.title,
                font=(self.font_family, 14, "bold"),
                bg=ModernColors.CARD_BG,
                fg=ModernColors.TEXT_PRIMARY).pack(side=tk.LEFT)
        
        # Status dot with sunset accent
        tk.Label(header_frame, text="●",
                font=("Arial", 8),
                fg=ModernColors.GRADIENT_3,  # Pink accent dot
                bg=ModernColors.CARD_BG).pack(side=tk.RIGHT)
    
    def add_content(self, content_widget):
        """Add content to the card"""
        content_widget.pack(fill=tk.BOTH, expand=True)

class ModernButton(tk.Button):
    """Modern styled button"""
    
    def __init__(self, parent, text, command, style="primary", **kwargs):
        self.font_family = self.get_system_font()
        
        # Style configurations with sunset theme
        styles = {
            "primary": {
                "bg": ModernColors.PRIMARY,         # Rich blue #355C7D
                "hover_bg": ModernColors.PRIMARY_HOVER,
                "fg": "white"
            },
            "success": {
                "bg": ModernColors.GRADIENT_3,      # Pink
                "hover_bg": "#B85A7A",              # Darker pink
                "fg": "white"
            },
            "secondary": {
                "bg": ModernColors.GLASS_BG,
                "hover_bg": ModernColors.GLASS_HOVER,
                "fg": ModernColors.TEXT_PRIMARY
            }
        }
        
        config = styles.get(style, styles["primary"])
        
        super().__init__(
            parent,
            text=text,
            command=command,
            font=(self.font_family, 12, "bold"),
            bg=config["bg"],
            fg=config["fg"],
            activebackground=config["bg"],
            activeforeground=config["fg"],
            bd=0,
            relief="flat",
            padx=24,
            pady=12,
            cursor="hand2",
            **kwargs
        )
        
        self.hover_bg = config["hover_bg"]
        self.normal_bg = config["bg"]
        
        self.bind("<Enter>", self.on_enter)
        self.bind("<Leave>", self.on_leave)
    
    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")
    
    def on_enter(self, event):
        self.configure(bg=self.hover_bg)
    
    def on_leave(self, event):
        self.configure(bg=self.normal_bg)

class ModernEntry(tk.Frame):
    """Modern styled entry with label"""
    
    def __init__(self, parent, label, variable=None, **kwargs):
        super().__init__(parent, bg=ModernColors.BG_PRIMARY)
        self.font_family = self.get_system_font()
        self.variable = variable
        
        # Label
        tk.Label(self, text=label,
                font=(self.font_family, 10, "bold"),
                bg=ModernColors.BG_PRIMARY,
                fg=ModernColors.TEXT_PRIMARY).pack(anchor="w", pady=(0, 3))
        
        # Entry
        entry_frame = tk.Frame(self, bg=ModernColors.CARD_BORDER, bd=1)
        entry_frame.pack(fill=tk.X)
        
        self.entry = tk.Entry(entry_frame,
                             textvariable=variable,
                             font=(self.font_family, 10),
                             bg=ModernColors.BG_SECONDARY,  # Much darker background
                             fg=ModernColors.TEXT_PRIMARY,
                             insertbackground=ModernColors.TEXT_PRIMARY,
                             relief="flat",
                             bd=0,
                             **kwargs)
        self.entry.pack(fill=tk.X, ipady=6, ipadx=10)
    
    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")
    
    def get(self):
        return self.entry.get() if not self.variable else self.variable.get()

class ModernCombobox(tk.Frame):
    """Modern styled combobox with label"""
    
    def __init__(self, parent, label, variable=None, values=None, **kwargs):
        super().__init__(parent, bg=ModernColors.BG_PRIMARY)
        self.font_family = self.get_system_font()
        
        # Labelt
        tk.Label(self, text=label,
                font=(self.font_family, 10, "bold"),
                bg=ModernColors.BG_PRIMARY,
                fg=ModernColors.TEXT_PRIMARY).pack(anchor="w", pady=(0, 3))
        
        # Get the global style
        style = ttk.Style()
        
        # Force override all combobox styling
        style.configure("TCombobox",
                       fieldbackground=ModernColors.BG_SECONDARY,
                       background=ModernColors.BG_SECONDARY,
                       foreground=ModernColors.TEXT_PRIMARY,
                       borderwidth=1,
                       relief="flat",
                       selectbackground=ModernColors.PRIMARY,
                       selectforeground="white",
                       arrowcolor=ModernColors.TEXT_PRIMARY,
                       insertcolor=ModernColors.TEXT_PRIMARY,
                       lightcolor=ModernColors.BG_SECONDARY,
                       darkcolor=ModernColors.BG_SECONDARY,
                       bordercolor=ModernColors.CARD_BORDER,
                       focuscolor=ModernColors.PRIMARY)
        
        # Aggressive state mapping to override everything
        style.map("TCombobox",
                 fieldbackground=[("readonly", ModernColors.BG_SECONDARY),
                                ("active", ModernColors.BG_SECONDARY),
                                ("focus", ModernColors.BG_SECONDARY),
                                ("disabled", ModernColors.BG_SECONDARY),
                                ("pressed", ModernColors.BG_SECONDARY),
                                ("!readonly", ModernColors.BG_SECONDARY),
                                ("!focus", ModernColors.BG_SECONDARY),
                                ("!active", ModernColors.BG_SECONDARY)],
                 background=[("readonly", ModernColors.BG_SECONDARY),
                           ("active", ModernColors.BG_SECONDARY),
                           ("focus", ModernColors.BG_SECONDARY),
                           ("disabled", ModernColors.BG_SECONDARY),
                           ("pressed", ModernColors.BG_SECONDARY),
                           ("!readonly", ModernColors.BG_SECONDARY),
                           ("!focus", ModernColors.BG_SECONDARY),
                           ("!active", ModernColors.BG_SECONDARY)],
                 foreground=[("readonly", ModernColors.TEXT_PRIMARY),
                           ("active", ModernColors.TEXT_PRIMARY),
                           ("focus", ModernColors.TEXT_PRIMARY),
                           ("disabled", ModernColors.TEXT_MUTED),
                           ("pressed", ModernColors.TEXT_PRIMARY),
                           ("!readonly", ModernColors.TEXT_PRIMARY)],
                 bordercolor=[("focus", ModernColors.PRIMARY),
                            ("active", ModernColors.CARD_BORDER),
                            ("readonly", ModernColors.CARD_BORDER),
                            ("!focus", ModernColors.CARD_BORDER)])
        
        self.combobox = ttk.Combobox(self,
                                    textvariable=variable,
                                    values=values or [],
                                    font=(self.font_family, 10),
                                    style="TCombobox",  # Use the global style
                                    state="readonly",
                                    **kwargs)
        self.combobox.pack(fill=tk.X, ipady=3)
        
        # Force update and apply theme immediately
        self.combobox.update()
        self.update_idletasks()
    
    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")

class ModernProgressBar(tk.Frame):
    """Modern progress bar with label"""
    
    def __init__(self, parent):
        super().__init__(parent, bg=ModernColors.BG_PRIMARY)
        self.font_family = self.get_system_font()
        
        # Progress label with fixed width to prevent shaking
        self.progress_label = tk.Label(self, text="Ready to scan",
                                      font=(self.font_family, 11, "bold"),
                                      bg=ModernColors.BG_PRIMARY,
                                      fg=ModernColors.TEXT_PRIMARY,
                                      width=50,  # Fixed width to prevent layout shifts
                                      anchor="w")  # Left align text
        self.progress_label.pack(anchor="w", pady=(0, 6))
        
        # Progress bar
        style = ttk.Style()
        style.configure("Modern.Horizontal.TProgressbar",
                       background=ModernColors.PRIMARY,
                       troughcolor=ModernColors.CARD_BG,
                       borderwidth=0,
                       lightcolor=ModernColors.PRIMARY,
                       darkcolor=ModernColors.PRIMARY)
        
        self.progress_bar = ttk.Progressbar(self,
                                           style="Modern.Horizontal.TProgressbar",
                                           mode="determinate",
                                           maximum=100)  # Set fixed maximum
        self.progress_bar.pack(fill=tk.X, ipady=6)
    
    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")
    
    def set_progress(self, value, text=None):
        self.progress_bar["value"] = value
        if text:
            # Truncate text if too long to prevent overflow
            if len(text) > 48:
                text = text[:45] + "..."
            self.progress_label.configure(text=text)

class ModernTextArea(tk.Frame):
    """Modern text area with scrollbar"""
    
    def __init__(self, parent, height=10):
        super().__init__(parent, bg=ModernColors.BG_PRIMARY)
        self.font_family = self.get_system_font()
        
        # Create text widget with scrollbar
        text_frame = tk.Frame(self, bg=ModernColors.CARD_BG)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        self.text_widget = tk.Text(text_frame,
                                  height=height,
                                  font=(self.font_family, 10),
                                  bg=ModernColors.CARD_BG,
                                  fg=ModernColors.TEXT_PRIMARY,
                                  insertbackground=ModernColors.TEXT_PRIMARY,
                                  relief="flat",
                                  bd=0,
                                  wrap=tk.WORD,
                                  padx=12,
                                  pady=8)
        
        scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=self.text_widget.yview)
        self.text_widget.configure(yscrollcommand=scrollbar.set)
        
        self.text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    
    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")
    
    def insert(self, text):
        self.text_widget.insert(tk.END, text)
        self.text_widget.see(tk.END)
    
    def clear(self):
        self.text_widget.delete(1.0, tk.END)

class DocXScanApp:
    """Modern DocXScan application"""
    
    def __init__(self):
        self.root = None
        self.font_family = self.get_system_font()
        
        # State variables
        self.token_map = {}
        self.config_file = "docxscan.ini"
        
        # UI variables - will be initialized after root window is created
        self.selected_folder = None
        self.zip_folder = None
        self.zip_name = None
        self.file_type_choice = None
        self.custom_token_input = None
        self.loaded_token_file = None
        self.selected_token_label = None
        
        # Components
        self.progress_bar = None
        self.console = None
        self.token_dropdown = None
        
    def get_system_font(self):
        system = platform.system()
        fonts = {
            "Darwin": "SF Pro Display",
            "Windows": "Segoe UI",
            "Linux": "Ubuntu"
        }
        return fonts.get(system, "Arial")
    
    def create_window(self):
        """Create the modern window"""
        if self.root:
            try:
                self.root.destroy()
            except:
                pass
        
        self.root = tk.Tk()
        self.root.title("DocXScan v3.0 - Professional Document Scanner")
        self.root.geometry("1300x800")
        self.root.configure(bg=ModernColors.BG_PRIMARY)
        self.root.resizable(True, True)
        self.root.minsize(1100, 700)
        
        # Configure global ttk theme for dark mode
        self.configure_dark_theme()
        
        # Initialize tkinter variables after root window is created
        self.initialize_variables()
        
        self.create_interface()        
        self.load_config()
        self.center_window()
    
    def configure_dark_theme(self):
        """Configure dark theme for all ttk widgets"""
        style = ttk.Style()
        
        # Force a dark theme
        try:
            style.theme_use('clam')  # Use clam as base for better customization
        except:
            pass
        
        # Configure global options for better dark theme support
        self.root.option_add('*TCombobox*Listbox.background', ModernColors.BG_SECONDARY)
        self.root.option_add('*TCombobox*Listbox.foreground', ModernColors.TEXT_PRIMARY)
        self.root.option_add('*TCombobox*Listbox.selectBackground', ModernColors.PRIMARY)
        self.root.option_add('*TCombobox*Listbox.selectForeground', 'white')
        self.root.option_add('*TCombobox*Listbox.borderWidth', '0')
        self.root.option_add('*TCombobox*Listbox.relief', 'flat')
        
        # Configure ALL TCombobox styles aggressively
        style.configure("TCombobox",
                       fieldbackground=ModernColors.BG_SECONDARY,
                       background=ModernColors.BG_SECONDARY,
                       foreground=ModernColors.TEXT_PRIMARY,
                       borderwidth=1,
                       relief="flat",
                       selectbackground=ModernColors.PRIMARY,
                       selectforeground="white",
                       arrowcolor=ModernColors.TEXT_PRIMARY,
                       insertcolor=ModernColors.TEXT_PRIMARY,
                       lightcolor=ModernColors.BG_SECONDARY,
                       darkcolor=ModernColors.BG_SECONDARY,
                       bordercolor=ModernColors.CARD_BORDER,
                       focuscolor=ModernColors.PRIMARY)
        
        # Map ALL possible states for TCombobox
        style.map("TCombobox",
                 fieldbackground=[("readonly", ModernColors.BG_SECONDARY),
                                ("active", ModernColors.BG_SECONDARY),
                                ("focus", ModernColors.BG_SECONDARY),
                                ("disabled", ModernColors.BG_SECONDARY),
                                ("pressed", ModernColors.BG_SECONDARY),
                                ("!readonly", ModernColors.BG_SECONDARY)],
                 background=[("readonly", ModernColors.BG_SECONDARY),
                           ("active", ModernColors.BG_SECONDARY),
                           ("focus", ModernColors.BG_SECONDARY),
                           ("disabled", ModernColors.BG_SECONDARY),
                           ("pressed", ModernColors.BG_SECONDARY),
                           ("!readonly", ModernColors.BG_SECONDARY)],
                 foreground=[("readonly", ModernColors.TEXT_PRIMARY),
                           ("active", ModernColors.TEXT_PRIMARY),
                           ("focus", ModernColors.TEXT_PRIMARY),
                           ("disabled", ModernColors.TEXT_MUTED),
                           ("pressed", ModernColors.TEXT_PRIMARY),
                           ("!readonly", ModernColors.TEXT_PRIMARY)],
                 bordercolor=[("focus", ModernColors.PRIMARY),
                            ("active", ModernColors.PRIMARY),
                            ("readonly", ModernColors.CARD_BORDER),
                            ("disabled", ModernColors.CARD_BORDER),
                            ("pressed", ModernColors.PRIMARY),
                            ("!readonly", ModernColors.CARD_BORDER)],
                 arrowcolor=[("readonly", ModernColors.TEXT_PRIMARY),
                           ("active", ModernColors.TEXT_PRIMARY),
                           ("focus", ModernColors.TEXT_PRIMARY),
                           ("disabled", ModernColors.TEXT_MUTED),
                           ("pressed", ModernColors.TEXT_PRIMARY),
                           ("!readonly", ModernColors.TEXT_PRIMARY)])
        
        # Configure scrollbar for dark theme
        style.configure("Vertical.TScrollbar",
                       background=ModernColors.CARD_BG,
                       troughcolor=ModernColors.BG_SECONDARY,
                       bordercolor=ModernColors.CARD_BORDER,
                       arrowcolor=ModernColors.TEXT_PRIMARY,
                       darkcolor=ModernColors.CARD_BG,
                       lightcolor=ModernColors.CARD_BG)
        
        # Map scrollbar states
        style.map("Vertical.TScrollbar",
                 background=[("active", ModernColors.CARD_HOVER),
                           ("pressed", ModernColors.PRIMARY)])
        
        # Configure progress bar theme with sunset colors
        style.configure("Modern.Horizontal.TProgressbar",
                       background=ModernColors.GRADIENT_3,  # Pink progress bar
                       troughcolor=ModernColors.CARD_BG,
                       borderwidth=0,
                       lightcolor=ModernColors.GRADIENT_3,
                       darkcolor=ModernColors.GRADIENT_3)
        
        # Also configure any other potential ttk widgets
        style.configure("TEntry",
                       fieldbackground=ModernColors.BG_SECONDARY,
                       background=ModernColors.BG_SECONDARY,
                       foreground=ModernColors.TEXT_PRIMARY,
                       bordercolor=ModernColors.CARD_BORDER,
                       insertcolor=ModernColors.TEXT_PRIMARY)
        
        style.map("TEntry",
                 fieldbackground=[("focus", ModernColors.BG_SECONDARY),
                                ("active", ModernColors.BG_SECONDARY)],
                 bordercolor=[("focus", ModernColors.PRIMARY),
                            ("active", ModernColors.PRIMARY)])


    def initialize_variables(self):
        """Initialize tkinter variables after root window exists"""
        self.selected_folder = tk.StringVar()
        self.zip_folder = tk.StringVar()
        self.zip_name = tk.StringVar(value="matched_files")
        self.file_type_choice = tk.StringVar(value="Both (.docx and .dcp.docx)")
        self.custom_token_input = tk.StringVar()
        self.loaded_token_file = tk.StringVar(value="No token file loaded")
        self.selected_token_label = tk.StringVar(value="-- Select Token --")
    
    def create_interface(self):
        """Create the modern interface"""
        # Background
        bg_canvas = BlurredBackground(self.root, 1300, 800)
        bg_canvas.pack(fill=tk.BOTH, expand=True)
        
        # Main overlay
        main_overlay = tk.Frame(bg_canvas, bg=ModernColors.BG_PRIMARY)
        main_overlay.place(relx=0, rely=0, relwidth=1, relheight=1)
        
        self.create_header(main_overlay)
        self.create_main_content(main_overlay)
        self.create_footer(main_overlay)
    
    def create_header(self, parent):
        """Create header section"""
        header_frame = tk.Frame(parent, bg=ModernColors.BG_PRIMARY, height=70)
        header_frame.pack(fill=tk.X, padx=30, pady=(20, 10))
        header_frame.pack_propagate(False)
        
        # Title
        tk.Label(header_frame,
                text="DocXScan v3.0",
                font=(self.font_family, 24, "bold"),
                bg=ModernColors.BG_PRIMARY,
                fg=ModernColors.TEXT_PRIMARY).pack(anchor="w")
        
        tk.Label(header_frame,
                text="Professional document scanner with intelligent token detection",
                font=(self.font_family, 12, "normal"),
                bg=ModernColors.BG_PRIMARY,
                fg=ModernColors.TEXT_SECONDARY).pack(anchor="w", pady=(2, 0))
    
    def create_main_content(self, parent):
        """Create main content area - compact layout"""
        content_frame = tk.Frame(parent, bg=ModernColors.BG_PRIMARY)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=(0, 10))
        
        # Left column - Configuration (60% width)
        left_frame = tk.Frame(content_frame, bg=ModernColors.BG_PRIMARY)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 15))
        
        # Right column - Results (40% width)
        right_frame = tk.Frame(content_frame, bg=ModernColors.BG_PRIMARY)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        self.create_config_section(left_frame)
        self.create_results_section(right_frame)
    
    def create_config_section(self, parent):
        """Create configuration section - compact"""
        # File Selection Card - More compact
        file_card = ModernCard(parent, "📁  File Selection")
        file_card.pack(fill=tk.X, pady=(0, 8))
        
        file_content = tk.Frame(file_card, bg=ModernColors.CARD_BG)
        file_content.pack(fill=tk.BOTH, expand=True, padx=12, pady=(0, 8))
        
        # Folder selection - more compact
        folder_frame = tk.Frame(file_content, bg=ModernColors.CARD_BG)
        folder_frame.pack(fill=tk.X, pady=(0, 6))
        
        tk.Label(folder_frame, text="Scan Folder:",
                font=(self.font_family, 10, "bold"),
                bg=ModernColors.CARD_BG,
                fg=ModernColors.TEXT_PRIMARY).pack(anchor="w", pady=(0, 2))
        
        browse_frame = tk.Frame(folder_frame, bg=ModernColors.CARD_BG)
        browse_frame.pack(fill=tk.X)
        
        # Smaller button
        browse_btn = ModernButton(browse_frame, "Browse", self.browse_folder)
        browse_btn.configure(padx=12, pady=6, font=(self.font_family, 9, "bold"))
        browse_btn.pack(side=tk.LEFT)
        
        self.folder_label = tk.Label(browse_frame, text="No folder selected",
                                    font=(self.font_family, 8),
                                    bg=ModernColors.CARD_BG,
                                    fg=ModernColors.TEXT_TERTIARY,
                                    width=30,  # Fixed width to prevent shaking
                                    anchor="w")  # Left align text
        self.folder_label.pack(side=tk.LEFT, padx=(8, 0))
        
        # ZIP Output settings - more compact
        zip_folder_frame = tk.Frame(file_content, bg=ModernColors.CARD_BG)
        zip_folder_frame.pack(fill=tk.X, pady=(0, 6))
        
        tk.Label(zip_folder_frame, text="Output Folder:",
                font=(self.font_family, 10, "bold"),
                bg=ModernColors.CARD_BG,
                fg=ModernColors.TEXT_PRIMARY).pack(anchor="w", pady=(0, 2))
        
        zip_browse_frame = tk.Frame(zip_folder_frame, bg=ModernColors.CARD_BG)
        zip_browse_frame.pack(fill=tk.X)
        
        zip_btn = ModernButton(zip_browse_frame, "Browse", self.browse_zip_folder, style="secondary")
        zip_btn.configure(padx=12, pady=6, font=(self.font_family, 9, "bold"))
        zip_btn.pack(side=tk.LEFT)
        
        self.zip_folder_label = tk.Label(zip_browse_frame, text="No output folder selected",
                                        font=(self.font_family, 8),
                                        bg=ModernColors.CARD_BG,
                                        fg=ModernColors.TEXT_TERTIARY,
                                        width=30,  # Fixed width to prevent shaking
                                        anchor="w")  # Left align text
        self.zip_folder_label.pack(side=tk.LEFT, padx=(8, 0))
        
        # Compact inputs in two columns to save space
        inputs_frame = tk.Frame(file_content, bg=ModernColors.CARD_BG)
        inputs_frame.pack(fill=tk.X, pady=(0, 6))
        
        # Left column
        left_inputs = tk.Frame(inputs_frame, bg=ModernColors.CARD_BG)
        left_inputs.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 4))
        
        zip_name_entry = ModernEntry(left_inputs, "ZIP Name:", self.zip_name)
        zip_name_entry.pack(fill=tk.X)
        
        # Right column
        right_inputs = tk.Frame(inputs_frame, bg=ModernColors.CARD_BG)
        right_inputs.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(4, 0))
        
        file_type_combo = ModernCombobox(right_inputs, "File Type:", self.file_type_choice,
                                        values=["Both (.docx and .dcp.docx)", "Only .dcp.docx", "Only .docx (excluding .dcp.docx)"])
        file_type_combo.pack(fill=tk.X)
        
        # Token Configuration Card - compact
        token_card = ModernCard(parent, "🔍  Token Configuration")
        token_card.pack(fill=tk.X, pady=(0, 8))
        
        token_content = tk.Frame(token_card, bg=ModernColors.CARD_BG)
        token_content.pack(fill=tk.BOTH, expand=True, padx=12, pady=(0, 8))
        
        # Token file loading - more compact
        token_file_frame = tk.Frame(token_content, bg=ModernColors.CARD_BG)
        token_file_frame.pack(fill=tk.X, pady=(0, 6))
        
        tk.Label(token_file_frame, text="Token File:",
                font=(self.font_family, 10, "bold"),
                bg=ModernColors.CARD_BG,
                fg=ModernColors.TEXT_PRIMARY).pack(anchor="w", pady=(0, 2))
        
        token_load_frame = tk.Frame(token_file_frame, bg=ModernColors.CARD_BG)
        token_load_frame.pack(fill=tk.X)
        
        token_btn = ModernButton(token_load_frame, "Load File", self.load_token_file)
        token_btn.configure(padx=12, pady=6, font=(self.font_family, 9, "bold"))
        token_btn.pack(side=tk.LEFT)
        
        self.token_file_label = tk.Label(token_load_frame, text="No token file loaded",
                                        font=(self.font_family, 8),
                                        bg=ModernColors.CARD_BG,
                                        fg=ModernColors.TEXT_TERTIARY,
                                        width=35,  # Fixed width to prevent shaking
                                        anchor="w")  # Left align text
        self.token_file_label.pack(side=tk.LEFT, padx=(8, 0))
        
        # Token selection and custom tokens in two columns
        token_inputs_frame = tk.Frame(token_content, bg=ModernColors.CARD_BG)
        token_inputs_frame.pack(fill=tk.X, pady=(0, 6))
        
        # Left column - Token selection
        left_token = tk.Frame(token_inputs_frame, bg=ModernColors.CARD_BG)
        left_token.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 4))
        
        self.token_dropdown = ModernCombobox(left_token, "Select Token:", self.selected_token_label,
                                            values=["-- Select Token --"])
        self.token_dropdown.pack(fill=tk.X)
        
        # Right column - Custom tokens
        right_token = tk.Frame(token_inputs_frame, bg=ModernColors.CARD_BG)
        right_token.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(4, 0))
        
        custom_tokens_entry = ModernEntry(right_token, "Custom Tokens:", self.custom_token_input)
        custom_tokens_entry.pack(fill=tk.X)
        
        # Controls Card - compact and expandable to fill remaining space
        controls_card = ModernCard(parent, "⚡  Scan Controls")
        controls_card.pack(fill=tk.BOTH, expand=True)  # Changed to expand to fill remaining space
        
        controls_content = tk.Frame(controls_card, bg=ModernColors.CARD_BG)
        controls_content.pack(fill=tk.BOTH, expand=True, padx=12, pady=(0, 8))
        
        # Progress bar - more compact
        self.progress_bar = ModernProgressBar(controls_content)
        self.progress_bar.pack(fill=tk.X, pady=(0, 6))
        
        # Control buttons - smaller and more compact
        button_frame = tk.Frame(controls_content, bg=ModernColors.CARD_BG)
        button_frame.pack(fill=tk.X)
        
        start_btn = ModernButton(button_frame, "🚀 Start Scan", self.run_scan_threaded)
        start_btn.configure(padx=14, pady=6, font=(self.font_family, 10, "bold"))
        start_btn.pack(side=tk.LEFT, padx=(0, 6))
        
        template_btn = ModernButton(button_frame, "📄 Template", self.create_template, style="secondary")
        template_btn.configure(padx=10, pady=6, font=(self.font_family, 9, "bold"))
        template_btn.pack(side=tk.LEFT, padx=(0, 6))
        
        clear_btn = ModernButton(button_frame, "🧹 Clear", self.clear_console, style="secondary")
        clear_btn.configure(padx=10, pady=6, font=(self.font_family, 9, "bold"))
        clear_btn.pack(side=tk.LEFT)
        
        # Add a spacer frame to push everything to the top and fill remaining space
        spacer_frame = tk.Frame(controls_content, bg=ModernColors.CARD_BG)
        spacer_frame.pack(fill=tk.BOTH, expand=True)
    
    def create_results_section(self, parent):
        """Create results section - compact"""
        # Console Card
        console_card = ModernCard(parent, "📋  Scan Results")
        console_card.pack(fill=tk.BOTH, expand=True)
        
        console_content = tk.Frame(console_card, bg=ModernColors.CARD_BG)
        console_content.pack(fill=tk.BOTH, expand=True, padx=12, pady=(0, 8))
        
        # Console area - optimized height
        self.console = ModernTextArea(console_content, height=16)
        self.console.pack(fill=tk.BOTH, expand=True)
    
    def create_footer(self, parent):
        """Create footer"""
        footer_frame = tk.Frame(parent, bg=ModernColors.BG_PRIMARY)
        footer_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=30, pady=(0, 15))
        
        tk.Label(footer_frame,
                text="© 2025 Hrishik Kunduru • DocXScan v3.0 Professional • All Rights Reserved",
                font=(self.font_family, 9, "normal"),
                bg=ModernColors.BG_PRIMARY,
                fg=ModernColors.TEXT_MUTED).pack()
    
    def center_window(self):
        """Center window on screen"""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')
    
    def log(self, message):
        """Log message to console"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        formatted_msg = f"[{timestamp}] {message}\n"
        self.console.insert(formatted_msg)
    
    def clear_console(self):
        """Clear console"""
        self.console.clear()
        self.log("Console cleared")
    
    def browse_folder(self):
        """Browse for scan folder"""
        path = filedialog.askdirectory(title="Select folder to scan")
        if path:
            self.selected_folder.set(path)
            # Truncate long folder names to prevent UI shaking
            display_name = os.path.basename(path)
            if len(display_name) > 28:
                display_name = display_name[:25] + "..."
            self.folder_label.configure(text=display_name)
            self.save_config()
            self.log(f"✅ Selected scan folder: {os.path.basename(path)}")
    
    def browse_zip_folder(self):
        """Browse for ZIP output folder"""
        path = filedialog.askdirectory(title="Select ZIP output folder")
        if path:
            self.zip_folder.set(path)
            # Truncate long folder names to prevent UI shaking
            display_name = os.path.basename(path)
            if len(display_name) > 28:
                display_name = display_name[:25] + "..."
            self.zip_folder_label.configure(text=display_name)
            self.save_config()
            self.log(f"✅ Selected output folder: {os.path.basename(path)}")
    
    def load_token_file(self):
        """Load token file"""
        path = filedialog.askopenfilename(
            title="Select Token JSON File",
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
        )
        if path:
            try:
                with open(path, "r", encoding='utf-8') as f:
                    self.token_map = json.load(f)
                
                # Truncate long filenames to prevent UI shaking
                filename = os.path.basename(path)
                if len(filename) > 20:
                    filename = filename[:17] + "..."
                
                display_text = f"Loaded: {filename} ({len(self.token_map)} tokens)"
                if len(display_text) > 33:
                    display_text = display_text[:30] + "..."
                
                self.loaded_token_file.set(display_text)
                self.token_file_label.configure(text=display_text)
                self.log(f"✅ Loaded {len(self.token_map)} tokens from {os.path.basename(path)}")
                
                # Update dropdown
                token_values = ["-- Select Token --"] + list(self.token_map.values())
                self.token_dropdown.combobox["values"] = token_values
                self.save_config()
                
            except Exception as e:
                error_text = "❌ Failed to load token file"
                self.loaded_token_file.set(error_text)
                self.token_file_label.configure(text=error_text)
                messagebox.showerror("Error", f"Failed to load token file:\n{e}")
                self.log(f"❌ Failed to load token file: {str(e)}")
    
    def create_template(self):
        """Create token template"""
        template = {
          "<<FileService.": "Fileservice",
          "</ff>": "Page Break",
          "</pp>": "Hard Return",
          "<backspace>": "Backspace",
          "<<STNDRDTH": "STNDRD Add \"TH\"",
          "<c>": "Center",
          "<u>": "Underline",
          "<i>": "Italic",
          "<pcase>": "pcase",
          "<lcase>": "lcase",
          "<ucase>": "ucase",
          "<bold>": "Bold",
          "<nobullet>": "No Bullet",
          "<fontsize": "Font Size",
          "<s1>": "s1",
          "<s2>": "s2",
          "[[MCOMPUTEINTO(<<": "MCOMPUTE INTO",
          "[[SCOMPUTEINTO(": "SCOMPUTE INTO",
          "[[ABORTIIF": "Abortif",
          "PROMTINTO(": "PROMTINTO",
          "PROMTINTOIIF(": "PROMTINTOIIF",
          "PROMTINTOLIST(": "PROMTINTOLIST",
          "PROMTINTOIIFLIST(": "PROMTINTOIIFLIST",
          "PROMTFORM(": "PROMTFORM",
          "<<Checklist.": "CHECKLIST",
          "TABLE(": "TABLE",
          "<<jfig": "JFIG",
          "jfig": "JFIG_General",
          "{ATTY": "ESIGN",
          "<<Special.": "SPECIAL",
          "+91|<<Special.ToDay": "+91 special day",
          "-91|<<Special.ToDay": "-91 special day",
          "+2|<<Special.ToDay": "+2 special day",
          "-2|<<Special.ToDay": "-2 special day",
          "<<Tracker.MortDate>>~MMMM dd": "MMMM dd,yyyy",
          "<<Tracker.MortDate>>~MM-dd-yyyy": "MM-dd-yyyy",
          "<<Tracker.MortDate>>~ddd": "ddd,MMM dd-yyyy",
          "<<Tracker.OriginalPrincipal>>~##": "##,###,###.00",
          "CU$TOMMMMMMPLACEHOLDER": "Custom(Enter Below)"
        }
        
        path = filedialog.asksaveasfilename(
            title="Save Token Template",
            defaultextension=".json",
            filetypes=[("JSON Files", "*.json")]
        )
        
        if path:
            try:
                with open(path, "w", encoding='utf-8') as f:
                    json.dump(template, f, indent=2)
                self.log(f"📄 Created template: {os.path.basename(path)}")
                messagebox.showinfo("Success", f"Template created successfully!\n{path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to create template:\n{e}")
                self.log(f"❌ Failed to create template: {str(e)}")
    
    def run_scan_threaded(self):
        """Start scan in thread"""
        self.progress_bar.set_progress(0, "Starting scan...")
        self.log("🔍 Starting document scan...")
        threading.Thread(target=self.scan_documents, daemon=True).start()
    
    def extract_full_text_lines(self, doc):
        """Extract text from document"""
        lines = []
        for para in doc.paragraphs:
            lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    lines.append(cell.text)
        return lines
    
    def scan_documents(self):
        """Main document scanning logic"""
        try:
            # Reset progress
            self.root.after(0, lambda: self.progress_bar.set_progress(0, "Initializing..."))
            
            folder = self.selected_folder.get().strip()
            zip_dest = self.zip_folder.get().strip()
            zip_filename_base = self.zip_name.get().strip()

            # Validation - 5% progress
            self.root.after(0, lambda: self.progress_bar.set_progress(5, "Validating inputs..."))
            
            if not folder or not os.path.isdir(folder):
                self.root.after(0, lambda: self.log("❌ Invalid folder selected"))
                self.root.after(0, lambda: self.progress_bar.set_progress(0, "Invalid folder"))
                return

            if not zip_dest or not os.path.isdir(zip_dest):
                self.root.after(0, lambda: self.log("❌ Invalid ZIP destination folder"))
                self.root.after(0, lambda: self.progress_bar.set_progress(0, "Invalid ZIP folder"))
                return

            if not zip_filename_base:
                self.root.after(0, lambda: self.log("❌ ZIP filename cannot be empty"))
                self.root.after(0, lambda: self.progress_bar.set_progress(0, "Invalid ZIP name"))
                return

            if not self.token_map:
                self.root.after(0, lambda: self.log("❌ No tokens loaded. Please load a token file first."))
                self.root.after(0, lambda: self.progress_bar.set_progress(0, "No tokens loaded"))
                return

            # Get selected token - 10% progress
            self.root.after(0, lambda: self.progress_bar.set_progress(10, "Processing tokens..."))
            
            selected_label = self.selected_token_label.get()
            matched_tokens = [k for k, v in self.token_map.items() if v == selected_label]
            if not matched_tokens and selected_label != "-- Select Token --":
                patterns = matched_tokens
            else:
                # If no specific token selected, ask user
                if selected_label == "-- Select Token --":
                    self.root.after(0, lambda: self.log("❌ Please select a valid token from the dropdown"))
                    self.root.after(0, lambda: self.progress_bar.set_progress(0, "No token selected"))
                    return
                patterns = matched_tokens

            # Add custom tokens
            custom_tokens = [t.strip() for t in self.custom_token_input.get().split(",") if t.strip()]
            for ct in custom_tokens:
                self.token_map[ct] = f"Custom: {ct}"
            patterns += custom_tokens

            if not patterns:
                self.root.after(0, lambda: self.log("❌ No patterns to search for. Please select a token or add custom tokens."))
                self.root.after(0, lambda: self.progress_bar.set_progress(0, "No patterns"))
                return

            # File type filter - 15% progress
            self.root.after(0, lambda: self.progress_bar.set_progress(15, "Setting up file filters..."))
            
            file_type_map = {
                "Only .dcp.docx": lambda f: f.endswith('.dcp.docx'),
                "Only .docx (excluding .dcp.docx)": lambda f: f.endswith('.docx') and not f.endswith('.dcp.docx'),
                "Both (.docx and .dcp.docx)": lambda f: f.endswith('.docx')
            }
            file_filter = file_type_map.get(self.file_type_choice.get())

            matching_files = []
            metadata = []

            # Scan files - 20% progress
            self.root.after(0, lambda: self.progress_bar.set_progress(20, "Collecting files..."))
            self.root.after(0, lambda: self.log(f"🔍 Scanning folder: {folder}"))
            self.root.after(0, lambda: self.log(f"📋 Looking for patterns: {', '.join(patterns[:3])}{'...' if len(patterns) > 3 else ''}"))
            
            all_files = []
            for root_dir, _, files in os.walk(folder):
                for file in files:
                    if file_filter(file) and not file.startswith('~'):
                        all_files.append(os.path.join(root_dir, file))

            if not all_files:
                self.root.after(0, lambda: self.log("❌ No files found to scan"))
                self.root.after(0, lambda: self.progress_bar.set_progress(0, "No files found"))
                return

            self.root.after(0, lambda: self.log(f"📄 Found {len(all_files)} files to scan"))
            self.root.after(0, lambda: self.progress_bar.set_progress(25, f"Found {len(all_files)} files"))

            # Process files - 25% to 80% progress
            scan_progress_start = 25
            scan_progress_range = 55  # 80% - 25% = 55%
            
            for i, full_path in enumerate(all_files):
                try:
                    # Calculate accurate progress for file scanning
                    file_progress = scan_progress_start + int((i / len(all_files)) * scan_progress_range)
                    progress_text = f"Scanning file {i + 1}/{len(all_files)}: {os.path.basename(full_path)[:30]}{'...' if len(os.path.basename(full_path)) > 30 else ''}"
                    
                    self.root.after(0, lambda p=file_progress, t=progress_text: 
                                   self.progress_bar.set_progress(p, t))
                    
                    doc = Document(full_path)
                    full_text = '\n'.join(self.extract_full_text_lines(doc))
                    matched = []
                    matched_lines = []

                    for token in patterns:
                        if token in full_text:
                            matched.append(token)
                            matched_lines.extend([line.strip() for line in full_text.split('\n') if token in line])

                    if matched:
                        matching_files.append(full_path)
                        info = os.stat(full_path)
                        token_count = sum(full_text.count(token) for token in matched)
                        
                        # Get token labels for matched tokens
                        token_labels = []
                        for token in matched:
                            if token in self.token_map:
                                token_labels.append(self.token_map[token])
                            else:
                                token_labels.append(token)
                        
                        metadata.append({
                            'File Name': os.path.basename(full_path),
                            'File Path': full_path,
                            'Size (bytes)': info.st_size,
                            'Creation Date': datetime.fromtimestamp(info.st_ctime).strftime('%Y-%m-%d %H:%M:%S'),
                            'Modified Date': datetime.fromtimestamp(info.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
                            'Matched Pattern(s)': ', '.join(token_labels),
                            'Matched Line(s)': '/----/'.join(matched_lines),
                            'Token Match Count': token_count,
                            'Link to File': f'=HYPERLINK("{full_path}", "Open File")'
                        })
                        
                        self.root.after(0, lambda fname=os.path.basename(full_path): 
                                       self.log(f"✅ Found match in: {fname}"))

                except Exception as e:
                    self.root.after(0, lambda path=full_path, error=str(e): 
                                   self.log(f"❌ Error processing {os.path.basename(path)}: {error}"))

            # Complete file scanning - 80% progress
            self.root.after(0, lambda: self.progress_bar.set_progress(80, "File scanning complete"))

            if not matching_files:
                self.root.after(0, lambda: self.log("ℹ️ No matching files found"))
                self.root.after(0, lambda: self.progress_bar.set_progress(100, "No matches found"))
                return

            # Create Excel file - 85% progress
            self.root.after(0, lambda: self.progress_bar.set_progress(85, "Creating Excel report..."))
            self.root.after(0, lambda: self.log("📊 Creating Excel report..."))
            excel_filename = os.path.join(folder, 'matching_files_metadata.xlsx')
            pd.DataFrame(metadata).to_excel(excel_filename, index=False)

            # Create ZIP archive - 90% progress
            self.root.after(0, lambda: self.progress_bar.set_progress(90, "Creating ZIP archive..."))
            self.root.after(0, lambda: self.log("🗜️ Creating ZIP archive..."))
            zip_path = os.path.join(zip_dest, zip_filename_base + '.zip')
            matched_folder = tempfile.mkdtemp(prefix='Matched_Files_')

            try:
                # Copy matching files - 95% progress
                self.root.after(0, lambda: self.progress_bar.set_progress(95, "Copying files to archive..."))
                for file in matching_files:
                    dest = os.path.join(matched_folder, os.path.basename(file))
                    shutil.copy2(file, dest)

                # Create ZIP - 98% progress
                self.root.after(0, lambda: self.progress_bar.set_progress(98, "Finalizing ZIP archive..."))
                with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    zipf.write(excel_filename, arcname='matching_files_metadata.xlsx')
                    for file in os.listdir(matched_folder):
                        zipf.write(os.path.join(matched_folder, file), 
                                 arcname=os.path.join('Matched_Files', file))

                # Complete - 100% progress
                self.root.after(0, lambda: self.progress_bar.set_progress(100, "Scan complete!"))
                self.root.after(0, lambda: self.log(f"✅ Scan complete! {len(matching_files)} matching files found"))
                self.root.after(0, lambda: self.log(f"📄 Excel report: {excel_filename}"))
                self.root.after(0, lambda: self.log(f"🗜️ ZIP archive: {zip_path}"))

            except Exception as e:
                self.root.after(0, lambda error=str(e): self.log(f"❌ Error creating output files: {error}"))
                self.root.after(0, lambda: self.progress_bar.set_progress(0, "Error creating files"))
            finally:
                # Cleanup temp folder
                try:
                    shutil.rmtree(matched_folder)
                except:
                    pass

        except Exception as e:
            self.root.after(0, lambda error=str(e): self.log(f"❌ Scan failed: {error}"))
            self.root.after(0, lambda: self.progress_bar.set_progress(0, "Scan failed"))
    
    def load_config(self):
        """Load configuration"""
        config = configparser.ConfigParser()
        if os.path.exists(self.config_file):
            try:
                config.read(self.config_file)
                self.zip_name.set(config.get("LAST", "zip_name", fallback="matched_files"))
                self.selected_folder.set(config.get("LAST", "scan_folder", fallback=""))
                self.zip_folder.set(config.get("LAST", "zip_folder", fallback=""))
                self.custom_token_input.set(config.get("LAST", "custom_tokens", fallback=""))
                
                # Update folder labels if paths exist
                if self.selected_folder.get():
                    display_name = os.path.basename(self.selected_folder.get())
                    if len(display_name) > 28:
                        display_name = display_name[:25] + "..."
                    self.folder_label.configure(text=display_name)
                if self.zip_folder.get():
                    display_name = os.path.basename(self.zip_folder.get())
                    if len(display_name) > 28:
                        display_name = display_name[:25] + "..."
                    self.zip_folder_label.configure(text=display_name)
                
                # Try to auto-load token file
                token_file = config.get("LAST", "token_file", fallback="")
                if token_file and os.path.exists(token_file):
                    try:
                        with open(token_file, "r", encoding='utf-8') as f:
                            self.token_map = json.load(f)
                        
                        # Truncate filename for display
                        filename = os.path.basename(token_file)
                        if len(filename) > 20:
                            filename = filename[:17] + "..."
                        
                        display_text = f"Auto-loaded: {filename} ({len(self.token_map)} tokens)"
                        if len(display_text) > 33:
                            display_text = display_text[:30] + "..."
                        
                        self.loaded_token_file.set(display_text)
                        self.token_file_label.configure(text=display_text)
                        token_values = ["-- Select Token --"] + list(self.token_map.values())
                        self.token_dropdown.combobox["values"] = token_values
                        self.log(f"✅ Auto-loaded {len(self.token_map)} tokens from config")
                    except Exception as e:
                        self.log(f"❌ Failed to auto-load token file: {str(e)}")
            except Exception as e:
                self.log(f"❌ Failed to load config: {str(e)}")
    
    def save_config(self):
        """Save configuration"""
        config = configparser.ConfigParser()
        config["LAST"] = {
            "scan_folder": self.selected_folder.get(),
            "zip_folder": self.zip_folder.get(),
            "zip_name": self.zip_name.get(),
            "custom_tokens": self.custom_token_input.get()
        }
        try:
            with open(self.config_file, "w") as f:
                config.write(f)
        except Exception as e:
            self.log(f"❌ Failed to save config: {str(e)}")
    
    def run(self):
        """Start the application"""
        try:
            self.create_window()
            self.root.protocol("WM_DELETE_WINDOW", self.root.destroy)
            self.log("🚀 DocXScan v3.0 Professional ready!")
            self.log("📝 Load a token file and select folders to begin scanning")
            self.root.mainloop()
        except Exception as e:
            print(f"Application error: {e}")

def main():
    """Main entry point"""
    app = DocXScanApp()
    app.run()

if __name__ == "__main__":
    main()


