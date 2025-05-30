import tkinter as tk
import tempfile
from tkinter import filedialog, messagebox, ttk
import os
import re
import zipfile
import pandas as pd
from datetime import datetime
from docx import Document
# -----------------------------------------------------------------------------
# DocXScan_1.5.py
# Copyright © 2025 Hrishik Kunduru. All rights reserved.
#
# Internal use only. Do not modify, redistribute, or reverse-engineer
# without written permission.
# -----------------------------------------------------------------------------

root = tk.Tk()
root.title("DocXScan v1.5")
root.geometry("650x700")

bg_color = "#1e1e1e"
fg_color = "#ffffff"
entry_bg = "#2d2d2d"
btn_bg = "#3c3c3c"
btn_fg = "#ffffff"

root.configure(bg=bg_color)

selected_folder = tk.StringVar()
zip_folder = tk.StringVar()
zip_name = tk.StringVar(value="matched_files")
file_type_choice = tk.StringVar(value="Both (.docx and .dcp.docx)")
pattern_choice = tk.StringVar(value="PROMTINTO")
custom_pattern = tk.StringVar()

style = ttk.Style()
style.theme_use("default")
style.configure("TCombobox", fieldbackground=entry_bg, background=entry_bg, foreground=fg_color)

def browse_folder():
    path = filedialog.askdirectory()
    if path:
        selected_folder.set(path)

def browse_zip_folder():
    path = filedialog.askdirectory()
    if path:
        zip_folder.set(path)

def log(msg):
    console_box.insert(tk.END, msg + "\n")
    console_box.see(tk.END)

def extract_full_text_lines(doc):
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.append(cell.text)
    return lines

def get_matching_lines_combined(file_path, patterns):
    try:
        doc = Document(file_path)
        lines = extract_full_text_lines(doc)
        matched_lines = []
        matched_patterns = set()
        for line in lines:
            for pattern in patterns:
                if re.search(pattern, line):
                    matched_lines.append(line.strip())
                    matched_patterns.add(pattern)
        return list(matched_patterns), matched_lines
    except Exception as e:
        log(f"Error reading {file_path}: {e}")
        return [], []

# Progress bar widget
progress = ttk.Progressbar(root, orient="horizontal", length=600, mode="determinate")
progress.pack(padx=10, pady=(0, 10))

def scan_docs():
    folder = selected_folder.get().strip()
    zip_dest = zip_folder.get().strip()
    zip_filename_base = zip_name.get().strip()

    if not folder or not os.path.isdir(folder):
        messagebox.showerror("Invalid folder", "Please select a valid folder to scan.")
        return
    if not zip_dest or not os.path.isdir(zip_dest):
        messagebox.showerror("Invalid ZIP folder", "Please select a valid ZIP destination.")
        return
    if not zip_filename_base:
        messagebox.showerror("Missing ZIP name", "Please enter a name for the ZIP file.")
        return

    file_type_map = {
        "Only .dcp.docx": lambda f: f.endswith('.dcp.docx'),
        "Only .docx (excluding .dcp.docx)": lambda f: f.endswith('.docx') and not f.endswith('.dcp.docx'),
        "Both (.docx and .dcp.docx)": lambda f: f.endswith('.docx')
    }

    file_filter = file_type_map.get(file_type_choice.get())
    if not file_filter:
        messagebox.showerror("Invalid selection", "Invalid file type choice.")
        return

    pattern_map = {
        "PROMTINTO": [r'PROMTINTO\('],
        "PROMTINTOIIF": [r'PROMTINTOIIF\('],
        "PROMTINTOLIST": [r'PROMTINTOLIST\('],
        "PROMTINTOIIFLIST": [r'PROMTINTOIIFLIST\('],
        "PROMTFORM": [r'PROMTFORM\('],
        "CHECKLIST": [r'<<Checklist.'],
        "TABLES": [r'TABLE\('],
        "JFIG": [r'<<jfig'],
        "JFIG_General": [r'jfig'],
        "ESIGN": [r'{ATTY'],
        "SPECIAL" : [R'<<Special.']
    }

    choice = pattern_choice.get()
    if choice == "Custom (Type in prompt below)":
        try:
            patterns = [p.strip() for p in custom_pattern.get().split(',') if p.strip()]
            for pat in patterns:
                re.compile(pat)
        except re.error as e:
            messagebox.showerror("Invalid Pattern", f"Error in custom regex pattern:\n{str(e)}")
            return
    else:
        patterns = pattern_map.get(choice, [])

    matching_files = []
    metadata = []

    all_files = []
    for root_dir, _, files in os.walk(folder):
        for file in files:
            if file_filter(file):
                all_files.append(os.path.join(root_dir, file))

    progress["maximum"] = len(all_files)
    progress["value"] = 0
    root.update_idletasks()

    for i, full_path in enumerate(all_files):
        matched_patterns, matched_lines = get_matching_lines_combined(full_path, patterns)
        if matched_lines:
            matching_files.append(full_path)
            info = os.stat(full_path)
            metadata.append({
                'File Name': os.path.basename(full_path),
                'Size (bytes)': info.st_size,
                'Creation Date': datetime.fromtimestamp(info.st_ctime).strftime('%Y-%m-%d %H:%M:%S'),
                'Modified Date': datetime.fromtimestamp(info.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
                'Matched Pattern': ', '.join(matched_patterns),
                'Matched Line(s)': '/----/'.join(matched_lines)
            })
        progress["value"] = i + 1
        root.update_idletasks()

    if not matching_files:
        log("No matching files found.")
        return

    excel_filename = os.path.join(folder, 'matching_files_metadata.xlsx')
    pd.DataFrame(metadata).to_excel(excel_filename, index=False)

    zip_path = os.path.join(zip_dest, zip_filename_base + '.zip')
    matched_folder = tempfile.mkdtemp(prefix='Matched_Files_')  # Temporary directory
    # Directory already created by tempfile
    for file in matching_files:
        dest = os.path.join(matched_folder, os.path.basename(file))
        with open(file, 'rb') as src, open(dest, 'wb') as dst:
            dst.write(src.read())

    with zipfile.ZipFile(zip_path, 'w') as zipf:
        zipf.write(excel_filename, arcname='matching_files_metadata.xlsx')
        for file in os.listdir(matched_folder):
            zipf.write(os.path.join(matched_folder, file), arcname=os.path.join('Matched_Files', file))

    log(f"\n✅ Done. {len(matching_files)} matching files found.")
    log(f"📄 Excel saved: {excel_filename}")
    log(f"🗜️ ZIP archive created: {zip_path}")

tk.Label(root, text="1. Select Folder to Scan:", bg=bg_color, fg=fg_color).pack(anchor="w", padx=10, pady=(10, 0))
tk.Button(root, text="Browse Folder", command=browse_folder, bg=btn_bg, fg=btn_fg).pack(anchor="w", padx=10)
tk.Label(root, textvariable=selected_folder, fg="#4aa3ff", bg=bg_color).pack(anchor="w", padx=10)

tk.Label(root, text="2. Select File Type:", bg=bg_color, fg=fg_color).pack(anchor="w", padx=10, pady=(10, 0))
ttk.Combobox(root, textvariable=file_type_choice, values=[
    "Only .dcp.docx",
    "Only .docx (excluding .dcp.docx)",
    "Both (.docx and .dcp.docx)"
], width=40).pack(anchor="w", padx=10)

tk.Label(root, text="3. Select Pattern:", bg=bg_color, fg=fg_color).pack(anchor="w", padx=10, pady=(10, 0))
ttk.Combobox(root, textvariable=pattern_choice, values=[
    "PROMTINTO",
    "PROMTINTOIIF",
    "PROMTINTOLIST",
    "PROMTINTOIIFLIST",
    "PROMTFORM",
    "CHECKLIST",
    "TABLES",
    "JFIG",
    "JFIG_General",
    "ESIGN",
    "SPECIAL",    
    "Custom (Type in prompt below)"
], width=40).pack(anchor="w", padx=10)

tk.Entry(root, textvariable=custom_pattern, width=80, bg=entry_bg, fg=fg_color, insertbackground=fg_color).pack(anchor="w", padx=10)

tk.Label(root, text="4. ZIP Destination Folder:", bg=bg_color, fg=fg_color).pack(anchor="w", padx=10, pady=(10, 0))
tk.Button(root, text="Browse ZIP Folder", command=browse_zip_folder, bg=btn_bg, fg=btn_fg).pack(anchor="w", padx=10)
tk.Label(root, textvariable=zip_folder, fg="#4aa3ff", bg=bg_color).pack(anchor="w", padx=10)

tk.Label(root, text="5. ZIP File Name:", bg=bg_color, fg=fg_color).pack(anchor="w", padx=10, pady=(10, 0))
tk.Entry(root, textvariable=zip_name, width=50, bg=entry_bg, fg=fg_color, insertbackground=fg_color).pack(anchor="w", padx=10)

tk.Button(root, text="Start Scan", bg="#007acc", fg="white", command=scan_docs).pack(pady=10)

tk.Label(root, text="Console Output:", bg=bg_color, fg=fg_color).pack(anchor="w", padx=10)
console_box = tk.Text(root, height=14, width=90, bg=entry_bg, fg=fg_color, insertbackground=fg_color)
console_box.pack(padx=10, pady=(0, 10))

tk.Label(root, text="© 2025 Hrishik Kunduru - All rights reserved",
         bg=bg_color, fg="#888888", font=("Arial", 9)).pack(side="bottom", pady=5)

root.mainloop()